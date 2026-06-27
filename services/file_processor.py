"""Обработка файлов и компрессия извлечённого текста (Маржа +95% Booster).

Фишка 2 мегаспеки NeuroMule 🐎⚡️: каждый текст, который мы извлекаем из .pdf /
.docx / .txt / .md / .csv / caption-сообщения, обязан пройти через
:func:`compress_extracted_text` перед склейкой в финальный промпт.

Для ``.pdf`` используем ``pypdf``. Сканы без текстового слоя рендерятся через
``pypdfium2`` → PNG для Vision в Нейротексте.

Лимиты размера документов:

* ``.xlsx`` / ``.xls`` / ``.csv`` — до **10 МБ**, скачивание чанками на диск
  в ``/tmp/`` (без удержания файла в RAM через ``BytesIO``);
* остальные форматы (``.pdf``, ``.docx``, ``.txt``, …) — до **15 МБ** в буфере.
"""

from __future__ import annotations

import base64
import logging
import math
import os
import re
import tempfile
from collections import Counter
from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path
from typing import Any, BinaryIO, Final

from services.table_number_parse import safe_float

logger = logging.getLogger(__name__)

# Видимые-под-микроскопом, но «токеноёмкие» символы: ZW-space, ZW-non-joiner,
# ZW-joiner, LRM / RLM, LRE / RLE / PDF / LRO / RLO, BOM, soft-hyphen, word-joiner.
_INVISIBLE_RE: Final = re.compile(
    r"[\u00ad\u200b-\u200f\u202a-\u202e\u2060\u2061\u2062\u2063\u2064\ufeff]"
)
_CARRIAGE_RE: Final = re.compile(r"\r\n?")
_TRAILING_WS_RE: Final = re.compile(r"[ \t]+(?=\n)")
_MULTI_SPACE_RE: Final = re.compile(r"[ \t]{2,}")
_MULTI_NEWLINE_RE: Final = re.compile(r"\n{3,}")

# Гард, чтобы не уронить ноду на 200-МБ "txt-бомбе".
DEFAULT_MAX_CHARS: Final = 500_000

# Жёсткий верхний лимит для табличных форматов (финансовые отчёты селлеров).
MAX_SPREADSHEET_BYTES: Final = 10 * 1024 * 1024  # 10 MB

# Лимит для прочих документов (.pdf, .docx, .txt, …).
MAX_DOCUMENT_BYTES: Final = 15 * 1024 * 1024  # 15 MB

SPREADSHEET_SUFFIXES: Final = frozenset({".xlsx", ".xls", ".csv"})

_DOWNLOAD_CHUNK_SIZE: Final = 64 * 1024

TXT_SPREADSHEET_TOO_BIG = (
    "⚠️ <b>Размер табличного файла превышает лимит 10 МБ.</b>\n\n"
    "Пожалуйста, сократите отчёт (меньше строк/листов) или отправьте сжатый "
    "файл .xlsx/.csv."
)

TXT_DOCUMENT_TOO_BIG = (
    "⚠️ <b>Размер файла превышает лимит 15 МБ.</b>\n\n"
    "Пожалуйста, сожмите документ и отправьте его снова. Подсказка: для "
    "PDF со сканами поможет повторное сохранение в PDF без OCR-слоя."
)


def _document_suffix(file_name: str | None) -> str:
    return Path((file_name or "document").strip()).suffix.lower()


def is_spreadsheet_suffix(suffix: str) -> bool:
    return (suffix or "").strip().lower() in SPREADSHEET_SUFFIXES


def max_document_bytes_for(file_name: str | None) -> int:
    """Возвращает лимит в байтах в зависимости от расширения файла."""
    if is_spreadsheet_suffix(_document_suffix(file_name)):
        return MAX_SPREADSHEET_BYTES
    return MAX_DOCUMENT_BYTES


def document_too_big_message(file_name: str | None) -> str:
    if is_spreadsheet_suffix(_document_suffix(file_name)):
        return TXT_SPREADSHEET_TOO_BIG
    return TXT_DOCUMENT_TOO_BIG


def _temp_download_dir() -> Path:
    preferred = Path("/tmp")
    if preferred.is_dir():
        return preferred
    return Path(tempfile.gettempdir())


def is_document_size_ok(
    size_bytes: int | None,
    *,
    file_name: str | None = None,
) -> bool:
    """Проверяет размер документа с учётом расширения (10 / 15 МБ).

    ``None`` → ``True`` (Telegram не всегда отдаёт точный размер).
    """

    if size_bytes is None:
        return True
    limit = max_document_bytes_for(file_name)
    return 0 <= int(size_bytes) <= limit


class DocumentTooBigError(ValueError):
    """Документ превышает лимит для своего типа."""

    def __init__(
        self,
        size_bytes: int,
        *,
        file_name: str | None = None,
        limit_bytes: int | None = None,
    ) -> None:
        limit = int(limit_bytes if limit_bytes is not None else max_document_bytes_for(file_name))
        super().__init__(f"document too big: {size_bytes} bytes > limit {limit}")
        self.size_bytes = int(size_bytes)
        self.limit_bytes = limit
        self.file_name = file_name


class _LimitingDiskWriter(BinaryIO):
    """Потоковая запись чанков на диск с жёстким лимитом размера."""

    def __init__(self, path: Path, *, max_bytes: int) -> None:
        self._path = path
        self._max_bytes = max_bytes
        self._written = 0
        self._fh = path.open("wb")
        self._closed = False

    def write(self, data: bytes) -> int:  # type: ignore[override]
        if self._closed:
            raise ValueError("writer is closed")
        nbytes = len(data)
        if self._written + nbytes > self._max_bytes:
            self._abort()
            raise DocumentTooBigError(
                self._written + nbytes,
                limit_bytes=self._max_bytes,
            )
        self._fh.write(data)
        self._written += nbytes
        return nbytes

    def _abort(self) -> None:
        self._closed = True
        try:
            self._fh.close()
        finally:
            self._path.unlink(missing_ok=True)

    def close(self) -> None:
        if not self._closed:
            self._fh.close()
            self._closed = True

    @property
    def size(self) -> int:
        return self._written

    def readable(self) -> bool:
        return False

    def writable(self) -> bool:
        return not self._closed

    def seekable(self) -> bool:
        return False

    def flush(self) -> None:
        self._fh.flush()


async def download_telegram_document_to_path(
    bot: Any,
    document: Any,
    *,
    file_name: str | None = None,
) -> str:
    """Скачивает табличный документ чанками в ``/tmp/``, возвращает путь к файлу.

    Только ``.xlsx``, ``.xls``, ``.csv`` — без удержания тела файла в RAM.
    Caller обязан удалить файл после обработки.
    """
    name = (file_name or getattr(document, "file_name", None) or "document").strip()
    suffix = _document_suffix(name)
    if not is_spreadsheet_suffix(suffix):
        raise ValueError(
            f"download_telegram_document_to_path supports spreadsheet files only, got {suffix!r}"
        )

    max_size = max_document_bytes_for(name)
    size_bytes = getattr(document, "file_size", None)
    if size_bytes is not None and not is_document_size_ok(size_bytes, file_name=name):
        raise DocumentTooBigError(int(size_bytes), file_name=name, limit_bytes=max_size)

    file_obj = await bot.get_file(document.file_id)
    tmp_dir = _temp_download_dir()
    fd, tmp_name = tempfile.mkstemp(prefix="neuromule_sheet_", suffix=suffix, dir=str(tmp_dir))
    os.close(fd)
    dest_path = Path(tmp_name)
    writer = _LimitingDiskWriter(dest_path, max_bytes=max_size)
    try:
        await bot.download_file(file_obj.file_path, destination=writer)
        writer.close()
        if not is_document_size_ok(writer.size, file_name=name):
            dest_path.unlink(missing_ok=True)
            raise DocumentTooBigError(writer.size, file_name=name, limit_bytes=max_size)
        return str(dest_path)
    except DocumentTooBigError:
        raise
    except Exception:
        dest_path.unlink(missing_ok=True)
        raise
    finally:
        if not writer._closed:
            writer.close()


async def download_telegram_document_to_buffer(
    bot: Any,
    document: Any,
    *,
    file_name: str | None = None,
    max_size: int | None = None,
) -> BytesIO:
    """Безопасно скачать Telegram ``Document`` в ``io.BytesIO`` (не для таблиц).

    Для ``.xlsx`` / ``.xls`` / ``.csv`` используйте
    :func:`download_telegram_document_to_path`.
    """
    name = (file_name or getattr(document, "file_name", None) or "document").strip()
    suffix = _document_suffix(name)
    if is_spreadsheet_suffix(suffix):
        raise ValueError(
            "spreadsheet documents must use download_telegram_document_to_path, not BytesIO"
        )

    limit = int(max_size if max_size is not None else max_document_bytes_for(name))
    size_bytes = getattr(document, "file_size", None)
    if size_bytes is not None and int(size_bytes) > limit:
        raise DocumentTooBigError(int(size_bytes), file_name=name, limit_bytes=limit)

    file_obj = await bot.get_file(document.file_id)
    buffer = BytesIO()
    await bot.download_file(file_obj.file_path, destination=buffer)
    actual = buffer.tell()
    if actual > limit:
        raise DocumentTooBigError(actual, file_name=name, limit_bytes=limit)
    buffer.seek(0)
    return buffer


def compress_extracted_text(raw_text: str) -> str:
    """Сжимает извлечённый из документа/текста сырой контент.

    Гарантии:

    * Удаляет невидимые / RTL-override / BOM символы.
    * Нормализует переводы строк (``\\r\\n`` → ``\\n``).
    * Стрипает пробелы в конце каждой строки.
    * ``[ \\t]{2,}`` → один пробел внутри строки.
    * ``\\n{3,}`` → ``\\n\\n`` (один пустой разделитель между абзацами).
    * Делает финальный ``strip``.

    Идемпотентна: повторный прогон через себя не меняет результат.
    """

    if not raw_text:
        return ""

    text = _CARRIAGE_RE.sub("\n", raw_text)
    text = _INVISIBLE_RE.sub("", text)
    text = _TRAILING_WS_RE.sub("", text)
    text = _MULTI_SPACE_RE.sub(" ", text)
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


# ─── Document extractors ───────────────────────────────────────────────────


def _safe_truncate(text: str, max_chars: int) -> str:
    if max_chars > 0 and len(text) > max_chars:
        return text[:max_chars]
    return text


def _read_txt(path: Path, max_chars: int) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    return _safe_truncate(raw, max_chars)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Извлекает текстовый слой PDF через ``pypdf`` (пустая строка = скан/картинки)."""
    try:
        import pypdf
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "pypdf не установлен: pip install 'pypdf>=4.0'"
        ) from exc

    reader = pypdf.PdfReader(BytesIO(file_bytes))
    text_layers: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_layers.append(page_text)
    return "\n".join(text_layers).strip()


def render_pdf_first_page_png(
    file_bytes: bytes,
    *,
    max_side_px: int = 1600,
) -> bytes | None:
    """Рендерит первую страницу PDF в PNG для Vision-OCR (сканы без текстового слоя)."""
    try:
        import pypdfium2 as pdfium
    except ImportError:
        logger.warning("pypdfium2 не установлен — Vision-fallback для PDF недоступен")
        return None

    try:
        pdf = pdfium.PdfDocument(file_bytes)
        if len(pdf) == 0:
            return None
        page = pdf[0]
        width, height = page.get_size()
        scale = min(
            max_side_px / max(width, 1.0),
            max_side_px / max(height, 1.0),
            3.0,
        )
        bitmap = page.render(scale=scale)
        pil_image = bitmap.to_pil()
        buffer = BytesIO()
        pil_image.save(buffer, format="PNG", optimize=True)
        return buffer.getvalue()
    except Exception:
        logger.exception("render_pdf_first_page_png failed")
        return None


def pdf_first_page_to_data_url(file_bytes: bytes) -> str | None:
    """Первая страница PDF → ``data:image/png;base64,...`` для OpenRouter Vision."""
    png = render_pdf_first_page_png(file_bytes)
    if not png:
        return None
    encoded = base64.standard_b64encode(png).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def _read_pdf(path: Path, max_chars: int) -> str:
    return _safe_truncate(extract_text_from_pdf(path.read_bytes()), max_chars)


def _read_docx(path: Path, max_chars: int) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx не установлен: добавьте 'python-docx>=1.1' в requirements.txt"
        ) from exc

    doc = Document(str(path))
    paragraphs: list[str] = []
    total = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if not text:
            continue
        paragraphs.append(text)
        total += len(text)
        if max_chars and total >= max_chars:
            break
    return _safe_truncate("\n".join(paragraphs), max_chars)


_WB_WAREHOUSE_ID_MAP: Final[dict[str, str]] = {
    "208547": "Рязань (Тюшевское)",
    "50003969": "Подольск (Транзит WB)",
}

_WB_AUX_AMOUNT_COL_HINTS: Final[tuple[str, ...]] = (
    "сумма",
    "всего",
    "total",
    "к удержанию",
    "удержано",
    "итого",
    "стоимост",
    "к оплате",
    "начислен",
    "размер",
    "перечисл",
    "тариф",
)
_WB_AUX_AMOUNT_HEADER_SKIP: Final[tuple[str, ...]] = (
    "обоснован",
    "описан",
    "тип док",
    "артикул",
    "бренд",
    "предмет",
    "наименован",
    "кол-во",
    "количество",
    "баркод",
    "штрих",
)
_WB_DETAIL_HEADER_HINTS: Final[tuple[str, ...]] = (
    "тип документа",
    "обоснован",
    "к перечислению",
    "артикул",
    "баркод",
)
_WB_WITHHOLDING_KEYWORDS: Final[tuple[str, ...]] = (
    "кредит",
    "выплат",
    "штраф",
    "удержан",
    "маркировк",
)
_WB_STORAGE_KEYWORDS: Final[tuple[str, ...]] = ("хранен",)


def map_wb_warehouse_label(value: str) -> str:
    """Замена цифровых ID складов WB на читаемые названия."""
    raw = (value or "").strip()
    if not raw:
        return raw
    key = raw.split(".")[0].strip()
    if key in _WB_WAREHOUSE_ID_MAP:
        return _WB_WAREHOUSE_ID_MAP[key]
    for wid, name in _WB_WAREHOUSE_ID_MAP.items():
        if wid in raw:
            return name
    return raw


def _wb_aux_sheet_category(sheet_name: str) -> str | None:
    """Категория вспомогательного листа WB: storage | system."""
    low = (sheet_name or "").lower().strip()
    if "хранен" in low or ("платн" in low and "хран" in low):
        return "storage"
    if "приемк" in low or "приёмк" in low:
        return "storage"
    if any(k in low for k in ("удержан", "кредит", "штраф", "санкц", "маркиров")):
        return "system"
    if "прочие" in low and any(k in low for k in ("удерж", "списан", "выплат")):
        return "system"
    return None


def _classify_aux_sheet_by_content(rows: list[list[str]]) -> str | None:
    """storage | system по тексту листа, если имя вкладки не распознано."""
    if not rows:
        return None
    sample = " ".join(
        str(cell or "") for row in rows[: min(25, len(rows))] for cell in row
    ).lower()
    if not sample.strip():
        return None
    if "хранен" in sample or ("платн" in sample and "хран" in sample):
        return "storage"
    if "приемк" in sample or "приёмк" in sample:
        return "storage"
    if any(
        k in sample
        for k in ("кредит", "удержан", "штраф", "санкц", "маркиров", "прочие удерж")
    ):
        return "system"
    return None


def _find_sheet_header_row(
    rows: list[list[str]],
    *,
    detail: bool = False,
    max_scan: int = 40,
) -> int:
    """Ищет строку шапки WB (после преамбулы поставщика на вспомогательных листах)."""
    hints = _WB_DETAIL_HEADER_HINTS if detail else _WB_AUX_AMOUNT_COL_HINTS
    for idx, row in enumerate(rows[:max_scan]):
        lowered = [_normalize_column_header(str(cell)) for cell in row]
        if not any(lowered):
            continue
        blob = " ".join(h for h in lowered if h)
        if detail:
            if any(k in blob for k in hints) and (
                "обоснован" in blob or "тип документа" in blob
            ):
                return idx
        elif any(k in blob for k in hints):
            return idx
    return 0


def _matrix_from_header_row(rows: list[list[str]], header_row: int) -> list[list[str]]:
    if header_row <= 0 or header_row >= len(rows):
        return rows
    return [list(rows[header_row]), *[list(r) for r in rows[header_row + 1 :]]]


def _is_plausible_wb_operation_amount(val: float) -> bool:
    """Одна операция WB: отсекает штрихкоды, ID и фантомные миллионы."""
    return 0.0 < val <= _WB_MAX_SINGLE_OPERATION_RUB


def _is_technical_sku_cell(text: str) -> bool:
    """Технический мусор шапки/итогов — не SKU для ABC-анализа."""
    raw = (text or "").strip().lower().replace("\u00a0", " ")
    if not raw:
        return True
    if any(word in raw for word in _SKU_TECHNICAL_JUNK_WORDS):
        return True
    compact = raw.replace(" ", "").replace(",", ".")
    if compact and re.fullmatch(r"[\d.]+", compact):
        return True
    return False


_WB_HEADER_SKU_MARKERS: Final[tuple[str, ...]] = ("баркод", "артикул", "barcode")
_WB_HEADER_LOGISTICS_MARKERS: Final[tuple[str, ...]] = (
    "логистик",
    "доставк",
    "delivery",
)
_WB_HEADER_CONTEXT_MARKERS: Final[tuple[str, ...]] = (
    "склад",
    "обоснован",
    "сумм",
)


def _wb_header_blob_contains_any(blob: str, markers: tuple[str, ...]) -> bool:
    return any(marker in blob for marker in markers)


def validate_wb_finance_detail_structure(matrix: list[list[str]] | None) -> bool:
    """
    Гибкая проверка шапки детализации WB: три группы неизменяемых e-commerce маркеров.

    Файл валиден, если в шапке (после преамбулы, до 40 строк) одновременно есть:
    артикул/баркод, логистика/доставка и склад/обоснование/сумма.
    Точные названия колонок различаются по месяцам — их подхватывает COLUMN_SYNONYMS.
    """
    if not matrix:
        return False
    for row in matrix[:40]:
        blob = " ".join(
            _normalize_column_header(str(cell)) for cell in row
        ).strip()
        if not blob:
            continue
        has_sku = _wb_header_blob_contains_any(blob, _WB_HEADER_SKU_MARKERS)
        has_logistics = _wb_header_blob_contains_any(blob, _WB_HEADER_LOGISTICS_MARKERS)
        has_context = _wb_header_blob_contains_any(blob, _WB_HEADER_CONTEXT_MARKERS)
        if has_sku and has_logistics and has_context:
            return True
    return False


def wb_finance_invalid_structure_payload(**extra: object) -> dict[str, object]:
    payload: dict[str, object] = {
        "error": WB_FINANCE_ERROR_INVALID_STRUCTURE,
        "cfo_build": _CFO_BUILD,
    }
    payload.update(extra)
    return payload


def check_wb_finance_upload_file(file_path: str) -> dict[str, object]:
    """Быстрая проверка структуры перед ETL (wb_ozon_finance / аудит WB)."""
    path = Path(file_path)
    if not path.is_file():
        return {"error": "empty_file", "cfo_build": _CFO_BUILD}
    if path.suffix.lower() == ".csv":
        matrix = _load_csv_matrix(path)
    else:
        loaded = load_cfo_workbook_from_path(file_path)
        matrix = loaded.matrix
    if not matrix or len(matrix) < 2:
        return {"error": "empty_file", "cfo_build": _CFO_BUILD}
    if not validate_wb_finance_detail_structure(matrix):
        return wb_finance_invalid_structure_payload()
    return {"ok": True, "cfo_build": _CFO_BUILD}


def _looks_like_rub_amount(text: str, val: float) -> bool:
    """Отсекает номера договоров и ID, оставляя денежные суммы."""
    raw = (text or "").strip().replace("\u00a0", " ")
    if not raw or not _is_plausible_wb_operation_amount(val):
        return False
    if re.search(r"\d[.,]\d{1,2}\b", raw.replace(" ", "")):
        return True
    if val >= 1_000_000:
        return False
    return val <= 999_999.99


def _row_money_fallback(row: list[str]) -> float:
    """Берёт наибольшую денежную ячейку строки, если колонка суммы не распознана."""
    best = 0.0
    for cell in row:
        text = str(cell or "").strip()
        if not text:
            continue
        val = abs(safe_float(cell))
        if _looks_like_rub_amount(text, val):
            best = max(best, val)
    return best


def _find_workbook_amount_columns(headers: list[str], *, allow_positional_fallback: bool = True) -> list[int]:
    """Индексы колонок с суммами на листах «Хранение» / «Удержания»."""
    lowered = [_normalize_column_header(str(h)) for h in headers]
    indices: list[int] = []
    for idx, header in enumerate(lowered):
        if not header:
            continue
        if any(skip in header for skip in _WB_AUX_AMOUNT_HEADER_SKIP):
            continue
        if any(hint in header for hint in _WB_AUX_AMOUNT_COL_HINTS):
            indices.append(idx)
    for idx, header in enumerate(lowered):
        if "услуг" in header and idx not in indices:
            indices.append(idx)
    if not indices:
        tx_cols = _resolve_cfo_tx_columns([str(h) for h in headers])
        if tx_cols is not None:
            for col in (tx_cols.payout_price, tx_cols.commission, tx_cols.retail_price):
                if col is not None and col not in indices:
                    indices.append(col)
    if not indices and allow_positional_fallback:
        for idx in range(len(headers) - 1, -1, -1):
            if str(headers[idx] or "").strip():
                indices.append(idx)
                break
    return indices


def _explicit_workbook_amount_columns(headers: list[str]) -> list[int]:
    lowered = [_normalize_column_header(str(h)) for h in headers]
    indices: list[int] = []
    for idx, header in enumerate(lowered):
        if header and any(hint in header for hint in _WB_AUX_AMOUNT_COL_HINTS):
            if any(skip in header for skip in _WB_AUX_AMOUNT_HEADER_SKIP):
                continue
            indices.append(idx)
    for idx, header in enumerate(lowered):
        if "услуг" in header and idx not in indices:
            indices.append(idx)
    return indices


def _row_workbook_amount(
    row: list[str],
    amount_cols: list[int],
    *,
    allow_money_fallback: bool = False,
) -> float:
    for col in amount_cols:
        if col < len(row):
            val = abs(safe_float(row[col]))
            if _is_plausible_wb_operation_amount(val):
                return val
    if allow_money_fallback:
        return _row_money_fallback(row)
    return 0.0


def _sum_workbook_sheet_amounts(rows: list[list[str]]) -> float:
    if len(rows) < 2:
        return 0.0
    header_row = _find_sheet_header_row(rows, detail=False)
    matrix = _matrix_from_header_row(rows, header_row)
    if len(matrix) < 2:
        return 0.0
    headers = [str(h) for h in matrix[0]]
    explicit_cols = _explicit_workbook_amount_columns(headers)
    amount_cols = explicit_cols or _find_workbook_amount_columns(headers)
    use_fallback = not explicit_cols
    total = sum(
        _row_workbook_amount(row, amount_cols, allow_money_fallback=use_fallback)
        for row in matrix[1:]
    )
    return _round_money(total)


def _should_skip_detail_withholding_row(blob: str, doc_type: str, *, kind: str) -> bool:
    """Строка «Кредит/Удержание» на детализации — дубль листа «Удержания»."""
    doc = (doc_type or "").strip().lower()
    text = f"{doc_type} {blob}".lower()
    if "удержан" in doc:
        return True
    if "кредит" in text:
        return True
    if kind == "system_loss" and any(k in text for k in ("кредит", "удержан")):
        return True
    return False


def _should_skip_detail_storage_row(blob: str, doc_type: str, *, kind: str) -> bool:
    """Строка хранения на детализации — дубль листа «Хранение»."""
    text = f"{doc_type} {blob}".lower()
    if kind in ("storage", "acceptance", "utilization"):
        return True
    return "хранен" in text


def scan_matrix_deep_costs(
    matrix: list[list[str]],
    *,
    skip_withholding_rows: bool = False,
    skip_storage_rows: bool = False,
) -> tuple[float, float]:
    """
    Сканирует лист детализации: хранение и системные удержания по обоснованию операции.
    """
    if not matrix or len(matrix) < 2:
        return 0.0, 0.0

    header_row = _find_sheet_header_row(matrix, detail=True)
    matrix = _matrix_from_header_row(matrix, header_row)
    if len(matrix) < 2:
        return 0.0, 0.0

    headers = [str(h) for h in matrix[0]]
    justification_cols: list[int] = []
    for idx, header in enumerate(headers):
        low = _normalize_column_header(header)
        if any(
            part in low
            for part in (
                "обоснован",
                "описание операции",
                "тип обоснования",
                "описание",
            )
        ):
            justification_cols.append(idx)

    amount_cols = _find_workbook_amount_columns(headers, allow_positional_fallback=False)
    tx_cols = _resolve_cfo_tx_columns(headers)
    if tx_cols is not None:
        if tx_cols.justification is not None and tx_cols.justification not in justification_cols:
            justification_cols.append(tx_cols.justification)
        if tx_cols.doc_type is not None and tx_cols.doc_type not in justification_cols:
            justification_cols.append(tx_cols.doc_type)
        for col in (tx_cols.payout_price, tx_cols.commission, tx_cols.retail_price):
            if col is not None and col not in amount_cols:
                amount_cols.append(col)

    storage_total = 0.0
    system_total = 0.0
    for row in matrix[1:]:
        if not any(str(cell or "").strip() for cell in row):
            continue
        blob = " ".join(
            str(row[col] or "") for col in justification_cols if col < len(row)
        ).lower()
        if not blob.strip():
            continue
        doc_type = ""
        if tx_cols is not None and tx_cols.doc_type is not None and tx_cols.doc_type < len(row):
            doc_type = str(row[tx_cols.doc_type] or "")
        row_kind = classify_cfo_tx_row(blob, doc_type=doc_type)
        if skip_withholding_rows and _should_skip_detail_withholding_row(
            blob, doc_type, kind=row_kind
        ):
            continue
        if skip_storage_rows and _should_skip_detail_storage_row(blob, doc_type, kind=row_kind):
            continue
        amount = _row_workbook_amount(row, amount_cols, allow_money_fallback=False)
        if amount <= 0 and (
            any(keyword in blob for keyword in _WB_STORAGE_KEYWORDS)
            or any(keyword in blob for keyword in _WB_WITHHOLDING_KEYWORDS)
        ):
            amount = _row_workbook_amount(row, amount_cols, allow_money_fallback=True)
        if amount <= 0:
            continue
        if any(keyword in blob for keyword in _WB_STORAGE_KEYWORDS):
            storage_total += amount
        elif any(keyword in blob for keyword in _WB_WITHHOLDING_KEYWORDS):
            system_total += amount

    return _round_money(storage_total), _round_money(system_total)


@dataclass(frozen=True)
class WbCfoAuxCostsContext:
    """Суммы с вспомогательных листов и флаги приоритета над детализацией."""

    storage: float = 0.0
    system: float = 0.0
    storage_from_dedicated_sheet: bool = False
    system_from_dedicated_sheet: bool = False


@dataclass
class CfoWorkbookLoadResult:
    """Результат глубокого чтения книги WB: детализация + вспомогательные листы."""

    matrix: list[list[str]]
    aux_storage_cost: float = 0.0
    aux_system_losses: float = 0.0
    storage_from_dedicated_sheet: bool = False
    system_from_dedicated_sheet: bool = False

    def aux_costs_context(self) -> WbCfoAuxCostsContext:
        return WbCfoAuxCostsContext(
            storage=self.aux_storage_cost,
            system=self.aux_system_losses,
            storage_from_dedicated_sheet=self.storage_from_dedicated_sheet,
            system_from_dedicated_sheet=self.system_from_dedicated_sheet,
        )


def _read_worksheet_rows(ws: object, *, max_rows: int = 50_000) -> list[list[str]]:
    rows: list[list[str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):  # type: ignore[attr-defined]
        if i >= max_rows:
            break
        cells = ["" if v is None else str(v).strip() for v in row]
        if any(cells):
            rows.append(cells)
    return rows


def _is_likely_wb_detail_sheet(rows: list[list[str]]) -> bool:
    if not rows:
        return False
    from services.wb_transaction_parse import is_wb_transaction_report

    header_row = _find_sheet_header_row(rows, detail=True)
    headers = [str(h) for h in rows[header_row]]
    if is_wb_transaction_report(headers):
        return True
    low = " ".join(_normalize_column_header(h) for h in headers)
    return "перечислен" in low and ("обоснован" in low or "артикул" in low)


def _detail_matrix_from_rows(rows: list[list[str]]) -> list[list[str]]:
    """Нормализует лист детализации: шапка после преамбулы WB."""
    if not rows:
        return rows
    header_row = _find_sheet_header_row(rows, detail=True)
    return _matrix_from_header_row(rows, header_row)


def load_cfo_workbook_from_path(file_path: str) -> CfoWorkbookLoadResult:
    """
    Глубокое чтение xlsx/csv: все листы, вспомогательные вкладки «Хранение» / «Удержания».
    """
    path = Path(file_path)
    if path.suffix.lower() == ".csv":
        matrix = _load_csv_matrix(path)
        detail_storage, detail_system = scan_matrix_deep_costs(matrix)
        return CfoWorkbookLoadResult(
            matrix=matrix,
            aux_storage_cost=detail_storage,
            aux_system_losses=detail_system,
        )

    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    detail_matrix: list[list[str]] | None = None
    sheet_storage = 0.0
    sheet_system = 0.0
    best_detail_score = -1

    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = _read_worksheet_rows(ws)
            if not rows:
                continue
            category = _wb_aux_sheet_category(sheet_name)
            if category is None and not _is_likely_wb_detail_sheet(rows):
                category = _classify_aux_sheet_by_content(rows)
            if category == "storage":
                sheet_storage += _sum_workbook_sheet_amounts(rows)
                continue
            if category == "system":
                sheet_system += _sum_workbook_sheet_amounts(rows)
                continue
            if _is_likely_wb_detail_sheet(rows):
                normalized = _detail_matrix_from_rows(rows)
                score = len(normalized[0]) * 1000 + len(normalized)
                if score > best_detail_score:
                    best_detail_score = score
                    detail_matrix = normalized
        if detail_matrix is None and wb.active is not None:
            detail_matrix = _detail_matrix_from_rows(_read_worksheet_rows(wb.active))
    finally:
        wb.close()

    storage_from_sheet = sheet_storage > 0
    system_from_sheet = sheet_system > 0
    detail_storage, detail_system = scan_matrix_deep_costs(
        detail_matrix or [],
        skip_withholding_rows=system_from_sheet,
        skip_storage_rows=storage_from_sheet,
    )
    if storage_from_sheet:
        aux_storage = _round_money(sheet_storage)
    else:
        aux_storage = _round_money(max(sheet_storage, detail_storage))
    if system_from_sheet:
        aux_system = _round_money(sheet_system)
    else:
        aux_system = _round_money(max(sheet_system, detail_system))

    return CfoWorkbookLoadResult(
        matrix=detail_matrix or [],
        aux_storage_cost=aux_storage,
        aux_system_losses=aux_system,
        storage_from_dedicated_sheet=storage_from_sheet,
        system_from_dedicated_sheet=system_from_sheet,
    )


def resolve_wb_cfo_workbook_input(
    *,
    file_path: str | Path | None = None,
    matrix_rows: list[list[str]] | None = None,
) -> tuple[list[list[str]], WbCfoAuxCostsContext]:
    """
    Матрица детализации WB + суммы с листов «Хранение» / «Удержания».

    При ``file_path`` читает книгу целиком; иначе сканирует переданную матрицу.
    """
    if file_path:
        loaded = load_cfo_workbook_from_path(str(file_path))
        matrix = loaded.matrix or list(matrix_rows or [])
        return matrix, loaded.aux_costs_context()
    rows = list(matrix_rows or [])
    storage, system = scan_matrix_deep_costs(rows)
    return rows, WbCfoAuxCostsContext(storage=storage, system=system)


def apply_deep_workbook_costs_to_engine(
    engine: CfoEngineResult,
    *,
    aux_storage_cost: float = 0.0,
    aux_system_losses: float = 0.0,
    aux_storage_from_sheet: bool = False,
    aux_system_from_sheet: bool = False,
) -> CfoEngineResult:
    """Подмешивает суммы с вспомогательных листов и пересчитывает чистую прибыль."""
    if aux_storage_from_sheet and aux_storage_cost > 0:
        storage = _round_money(aux_storage_cost + engine.total_storage_cost)
    else:
        storage = _round_money(max(engine.total_storage_cost, aux_storage_cost))
    if aux_system_from_sheet and aux_system_losses > 0:
        system = _round_money(aux_system_losses + engine.total_system_losses)
    else:
        system = _round_money(max(engine.total_system_losses, aux_system_losses))
    if (
        storage == engine.total_storage_cost
        and system == engine.total_system_losses
    ):
        return engine
    credit = max(
        engine.credit_deductions,
        aux_system_losses if aux_system_from_sheet and aux_system_losses > 0 else 0.0,
        system if aux_system_losses > 0 and not aux_system_from_sheet else 0.0,
    )
    clear_profit = _round_money(
        engine.total_sku_margin - storage - system - engine.tax_total
    )
    return replace(
        engine,
        total_storage_cost=storage,
        total_system_losses=system,
        credit_deductions=credit,
        clear_profit=clear_profit,
    )


def _load_csv_matrix(path: Path, *, max_rows: int = 5000) -> list[list[str]]:
    import csv

    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            with path.open("r", encoding=encoding, newline="") as fh:
                reader = csv.reader(fh)
                rows: list[list[str]] = []
                for i, row in enumerate(reader):
                    if i >= max_rows:
                        break
                    cells = [str(c).strip() for c in row]
                    if any(cells):
                        rows.append(cells)
                if rows:
                    return rows
        except UnicodeDecodeError:
            continue
    return []


def _read_xlsx_rows(path: Path, *, max_rows: int = 5000) -> list[list[str]]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows: list[list[str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= max_rows:
            break
        cells = ["" if v is None else str(v).strip() for v in row]
        if any(cells):
            rows.append(cells)
    wb.close()
    return rows


def read_xlsx_rows_from_path(path: Path | str, *, max_rows: int = 5000) -> list[list[str]]:
    """Читает строки Excel с диска (без загрузки всего файла в ``BytesIO``)."""
    return _read_xlsx_rows(Path(path), max_rows=max_rows)


def read_xlsx_rows_from_bytes(data: bytes, *, max_rows: int = 5000) -> list[list[str]]:
    """Читает Excel из небольшого in-memory буфера (только для тестов / микро-файлов)."""
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows: list[list[str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= max_rows:
            break
        cells = ["" if v is None else str(v).strip() for v in row]
        if any(cells):
            rows.append(cells)
    wb.close()
    return rows


def _read_xlsx(path: Path, max_chars: int) -> str:
    from services.table_markdown import rows_to_markdown_table

    rows = _read_xlsx_rows(path)
    if not rows:
        return ""
    md = rows_to_markdown_table(rows)
    return _safe_truncate(md, max_chars)


async def extract_text_from_document(
    path: Path | str,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    compress: bool = True,
) -> str:
    """Универсальный экстрактор: ``.txt/.md/.log/.csv`` + опц. ``.pdf/.docx``.

    Кидает ``ValueError`` для неизвестного расширения и ``RuntimeError`` если
    pypdf/python-docx не установлены. Возвращаемый текст уже пропущен через
    :func:`compress_extracted_text` (если ``compress=True``) — пайплайн чата
    может его склеивать с пользовательским caption без дополнительных шагов.
    """

    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in (".txt", ".md", ".log", ".csv"):
        raw = _read_txt(p, max_chars)
    elif suffix == ".pdf":
        raw = _read_pdf(p, max_chars)
    elif suffix in (".docx",):
        raw = _read_docx(p, max_chars)
    elif suffix == ".xlsx":
        raw = _read_xlsx(p, max_chars)
    else:
        raise ValueError(f"Unsupported document type: {suffix or '<no-ext>'}")

    if compress:
        return compress_extracted_text(raw)
    return raw


# ─── Локальный ETL товарной матрицы (ABC / FOMO логистики / OOS) ───────────

_USN_RATE = 0.06
_CFO_BUILD = "cfo-v12 (SaaS Protected Build)"
CFO_ENGINE_NAME = "CFO Engine v12 (Highload)"
WB_FINANCE_ERROR_INVALID_STRUCTURE = "invalid_structure"
_WB_MAX_SINGLE_OPERATION_RUB = 200_000.0
_SKU_TECHNICAL_JUNK_WORDS: Final[tuple[str, ...]] = (
    "выкупили",
    "акция",
    "дата",
    "номер",
    "итог",
    "шт",
)
_OOS_CRITICAL_DAYS = 5

# Семантический маппинг колонок WB (CFO Engine v11.1).
WB_COLUMN_SYNONYMS: dict[str, tuple[str, ...]] = {
    "sku": (
        "артикул",
        "арт",
        "артикул поставщика",
        "supplier_article",
        "sa_name",
        "barcode",
        "баркод",
        "штрихкод",
    ),
    "retail_price": (
        "продажа (ррц)",
        "цена розничная с учетом согласованной скидки",
        "розничная с согласованной",
        "ррц",
        "retail_price",
        "retail amount",
        "finishedprice",
        "доход",
        "выручка",
        "сумма продаж",
        "общая сумма продаж",
    ),
    "payout_price": (
        "к перечислению",
        "к перечислению за товар",
        "перечислению продавцу",
        "перечислению за",
        "сумма к перечислению",
        "вайлдберриз к перечислению",
        "wildberries_amount",
        "forpay",
        "payout_amount",
        "к начислению",
    ),
    "sale_price": (
        "продажа (ррц)",
        "цена розничная с учетом согласованной скидки",
        "розничная с согласованной",
        "ррц",
        "retail_price",
        "retail amount",
    ),
    "cost": ("себестоимость", "себес", "закупка", "cost", "supplier_price"),
    "delivery": (
        "логистика",
        "услуги по доставке",
        "доставка к клиенту",
        "delivery_amount",
    ),
    "return_delivery": (
        "доставка от клиента",
        "обратная логистика",
        "услуги по обратной доставке",
    ),
    "operation_type": (
        "тип документа",
        "обоснование для оплаты",
        "тип транзакции",
        "doc_type_name",
        "тип обоснования",
        "обоснован",
    ),
    "warehouse": (
        "наименование склада",
        "склад отгрузки",
        "логистический склад",
        "офис",
        "warehouse",
        "склад wb",
    ),
    "region": (
        "регион доставки",
        "регион",
        "область",
        "субъект",
        "населенный пункт",
        "город доставки",
        "region",
        "delivery region",
    ),
    "volume_liters": (
        "литраж",
        "объем",
        "объём",
        "volume",
        "литр",
        "габарит",
        "объем л",
    ),
}

# Публичные ключи как в интеграциях / ЛК WB (алиасы семантических полей).
COLUMN_SYNONYMS: dict[str, tuple[str, ...]] = {
    "sku": WB_COLUMN_SYNONYMS["sku"],
    "rrc_price": WB_COLUMN_SYNONYMS["retail_price"],
    "payout": WB_COLUMN_SYNONYMS["payout_price"],
    "delivery": WB_COLUMN_SYNONYMS["delivery"],
    "return_delivery": WB_COLUMN_SYNONYMS["return_delivery"],
    "op_type": WB_COLUMN_SYNONYMS["operation_type"],
    "warehouse": WB_COLUMN_SYNONYMS["warehouse"],
    "region": WB_COLUMN_SYNONYMS["region"],
}

_COLUMN_KEY_ALIASES: dict[str, str] = {
    "rrc_price": "retail_price",
    "payout": "payout_price",
    "op_type": "operation_type",
}

_WB_QTY_UNIT_HINTS = ("шт", "кол-во", "количество", "единиц")


def _normalize_column_header(header: str) -> str:
    return (header or "").lower().strip()


def find_column_index(
    df_columns: list[str],
    key_name: str,
    *,
    require_qty: bool = False,
    exclude_substrings: tuple[str, ...] = (),
) -> int | None:
    """
    Безопасный поиск колонки по ключу из :data:`WB_COLUMN_SYNONYMS`.

    Возвращает индекс или ``None`` (вызывающий код инициализирует нулями).
    """
    key_name = _COLUMN_KEY_ALIASES.get(key_name, key_name)
    synonyms = WB_COLUMN_SYNONYMS.get(key_name, ())
    if not synonyms:
        return None
    for idx, header in enumerate(df_columns):
        low = _normalize_column_header(header)
        if not low:
            continue
        if exclude_substrings and any(ex in low for ex in exclude_substrings):
            continue
        if require_qty and not any(q in low for q in _WB_QTY_UNIT_HINTS):
            continue
        if any(syn in low for syn in synonyms):
            if key_name == "return_delivery":
                if "возврат" in low or "обратн" in low or "от клиента" in low:
                    return idx
                continue
            if key_name in ("delivery", "sale_price", "cost") and (
                "возврат" in low or "обратн" in low
            ):
                continue
            return idx
    return None


def compute_buyout_coef_pct(sales_qty: float, returns_qty: float) -> float:
    """Выкуп: успешные продажи / (продажи + возвраты) × 100."""
    sales = max(0.0, float(sales_qty))
    returns = max(0.0, float(returns_qty))
    denom = sales + returns
    if denom > 0 and sales > 0:
        return sales / denom * 100.0
    if sales > 0:
        return 100.0
    return 0.0


def _round_money(value: float) -> float:
    return round(float(value), 2)


@dataclass
class CfoSkuBucket:
    """Поартикульный агрегат CFO Engine v11.1."""

    name: str
    article_id: str
    gross_sales_rrc: float = 0.0
    sales_qty: float = 0.0
    returns_qty: float = 0.0
    deliveries_qty: float = 0.0
    commission: float = 0.0
    forward_logistics: float = 0.0
    reverse_logistics: float = 0.0
    cost_rub: float = 0.0
    stock_qty: float = 0.0

    @property
    def revenue(self) -> float:
        return self.gross_sales_rrc

    @property
    def direct_expenses_rub(self) -> float:
        return self.commission + self.forward_logistics + self.reverse_logistics

    @property
    def margin_rub(self) -> float:
        return self.gross_sales_rrc - self.cost_rub - self.direct_expenses_rub

    @property
    def logistics(self) -> float:
        return self.forward_logistics + self.reverse_logistics

    @property
    def buyout_pct(self) -> float:
        return compute_buyout_coef_pct(self.sales_qty, self.returns_qty)


@dataclass
class CfoEngineResult:
    """Итог CFO Engine v11.1: налоговая база, корзины затрат, чистая прибыль."""

    kind: str
    tax_base_revenue: float
    tax_total: float
    total_sku_margin: float
    total_storage_cost: float
    total_system_losses: float
    total_ad_spend: float
    credit_deductions: float
    clear_profit: float
    sales_qty: float
    returns_qty: float
    buyout_coef_pct: float
    sku_buckets: dict[tuple[str, str], CfoSkuBucket]
    retail_price_source: str = "rrc"

    @property
    def commission_cost(self) -> float:
        return _round_money(sum(b.commission for b in self.sku_buckets.values()))

    @property
    def logistics_cost(self) -> float:
        return _round_money(
            sum(b.forward_logistics + b.reverse_logistics for b in self.sku_buckets.values())
        )

    @property
    def cost_of_goods(self) -> float:
        return _round_money(sum(b.cost_rub for b in self.sku_buckets.values()))

    @property
    def operational_profit(self) -> float:
        return _round_money(
            self.total_sku_margin - self.total_storage_cost - self.total_ad_spend - self.tax_total
        )


@dataclass(frozen=True)
class _CfoTxColumns:
    doc_type: int | None
    justification: int | None
    name: int | None
    article: int | None
    qty: int | None
    retail_price: int | None
    payout_price: int | None
    forward_logistics: int | None
    reverse_logistics: int | None
    commission: int | None
    cost: int | None
    stock: int | None


def _cfo_cell(row: list[str], col: int | None) -> str:
    if col is None or col >= len(row):
        return ""
    return str(row[col] or "").strip()


def _cfo_row_blob(row: list[str], cols: _CfoTxColumns) -> str:
    return " ".join(
        part
        for part in (
            _cfo_cell(row, cols.doc_type),
            _cfo_cell(row, cols.justification),
        )
        if part
    ).lower()


def classify_cfo_tx_row(blob: str, *, doc_type: str = "") -> str:
    """Классификация строки детализации WB для CFO Engine v11.1."""
    text = f"{doc_type} {blob}".lower().strip()
    if not text:
        return "skip"
    if "утилизац" in text:
        return "utilization"
    if "платная приемка" in text or "платная приёмка" in text:
        return "acceptance"
    if "хранен" in text or "стоимость хранения" in text:
        return "storage"
    if any(k in text for k in ("кредит", "взыскан", "честный знак", "маркировк")):
        return "system_loss"
    if "выплат" in text and any(k in text for k in ("кредит", "удержан", "штраф")):
        return "system_loss"
    if "штраф" in text:
        return "system_loss"
    if any(k in text for k in ("реклам", "продвижен", "трафарет", "спецразмещ", "медийн", "буст")):
        return "ad"
    if "сторно" in text:
        return "storno"
    if "отмена" in text:
        return "cancel"
    if "корректировк" in text and "продаж" in text:
        return "sale_adjustment"
    doc = (doc_type or "").lower().strip()
    if doc == "продажа" or text.startswith("продажа"):
        return "sale"
    if "возврат" in text and "логистик" not in text:
        return "return"
    if any(k in text for k in ("удержан", "штраф")) and "кредит" not in text:
        if any(k in text for k in ("реклам", "продвижен")):
            return "ad"
        return "system_loss"
    if any(k in text for k in ("логистик", "доставк", "перевоз")) and "хранен" not in text:
        if "обратн" in text or "от клиента" in text or ("возврат" in text and "логистик" in text):
            return "reverse_logistics"
        return "forward_logistics"
    if any(k in text for k in ("вознагражден", "комисс")):
        return "commission"
    return "other"


def _resolve_cfo_tx_columns(headers: list[str]) -> _CfoTxColumns | None:
    from services.wb_transaction_parse import is_wb_transaction_report

    if not is_wb_transaction_report(headers):
        return None
    lowered = [_normalize_column_header(h) for h in headers]

    def _find(*patterns: str, exclude: tuple[str, ...] = ()) -> int | None:
        for idx, header in enumerate(lowered):
            if not header:
                continue
            if exclude and any(ex in header for ex in exclude):
                continue
            if any(p in header for p in patterns):
                return idx
        return None

    doc_type = find_column_index(headers, "operation_type")
    if doc_type is None:
        doc_type = _find("тип документа")
    retail_price = find_column_index(headers, "retail_price")
    payout_price = find_column_index(headers, "payout_price")
    if payout_price is None:
        payout_price = _find("к перечислению", "перечислению продавцу", "перечислению за")
    reverse_logistics = find_column_index(headers, "return_delivery")
    forward_logistics = find_column_index(headers, "delivery", exclude_substrings=("хранен",))
    if forward_logistics is None:
        forward_logistics = _find("услуги по доставке", "логистик", "доставк", exclude=("хранен", "обратн", "возврат"))
    return _CfoTxColumns(
        doc_type=doc_type,
        justification=_find("обоснован"),
        name=_find("предмет", "наименование", "номенклатур", "бренд", "товар"),
        article=_find("артикул", "vendor", "sku", "barcode", "nmid", "штрих"),
        qty=_find("кол-во", "количество", "кол во", exclude=("возврат", "доставк", "заказ"))
        or _find("кол"),
        retail_price=retail_price,
        payout_price=payout_price,
        forward_logistics=forward_logistics,
        reverse_logistics=reverse_logistics,
        commission=_find("вознагражден", "комисс"),
        cost=find_column_index(headers, "cost"),
        stock=_find("остаток", "склад", "stock"),
    )


def _row_retail_amount(row: list[str], cols: _CfoTxColumns) -> tuple[float, str]:
    """Сумма продажи для налоговой базы: строго РРЦ, иначе fallback на перечисление."""
    if cols.retail_price is not None and cols.retail_price < len(row):
        val = safe_float(row[cols.retail_price])
        if val != 0:
            return val, "rrc"
    if cols.payout_price is not None and cols.payout_price < len(row):
        val = safe_float(row[cols.payout_price])
        if val != 0:
            return val, "payout_fallback"
    return 0.0, "missing"


def _row_direct_amount(row: list[str], col: int | None) -> float:
    if col is None or col >= len(row):
        return 0.0
    return abs(safe_float(row[col]))


def _cfo_sku_identity(row: list[str], cols: _CfoTxColumns) -> tuple[str, str]:
    name = _cfo_cell(row, cols.name) or _cfo_cell(row, cols.article) or "—"
    article = _cfo_cell(row, cols.article) or name
    if _is_technical_sku_cell(name) and _is_technical_sku_cell(article):
        return "—", "—"
    if _is_technical_sku_cell(name) and not _is_technical_sku_cell(article):
        name = article
    elif _is_technical_sku_cell(article) and not _is_technical_sku_cell(name):
        article = name
    return name[:64], article[:48]


def _cfo_get_bucket(
    buckets: dict[tuple[str, str], CfoSkuBucket],
    name: str,
    article: str,
) -> CfoSkuBucket:
    key = (name, article)
    bucket = buckets.get(key)
    if bucket is None:
        bucket = CfoSkuBucket(name=name, article_id=article)
        buckets[key] = bucket
    return bucket


def _finalize_cfo_engine(
    *,
    kind: str,
    buckets: dict[tuple[str, str], CfoSkuBucket],
    tax_base_revenue: float,
    total_storage_cost: float,
    total_system_losses: float,
    total_ad_spend: float,
    credit_deductions: float,
    sales_qty: float,
    returns_qty: float,
    retail_price_source: str,
    tax_preset: object | None = None,
) -> CfoEngineResult:
    from services.audit_tax import compute_audit_tax_total, default_wb_audit_tax_preset
    from services.wb_transaction_parse import is_valid_wb_sku

    preset = tax_preset or default_wb_audit_tax_preset()
    buckets = {
        key: bucket
        for key, bucket in buckets.items()
        if is_valid_wb_sku(bucket.name, bucket.article_id)
    }
    revenue_rrc = _round_money(max(0.0, tax_base_revenue))
    total_sku_margin = _round_money(sum(b.margin_rub for b in buckets.values()))
    _tax_base, tax_total = compute_audit_tax_total(
        preset=preset,
        tax_base_revenue=revenue_rrc,
        total_sku_margin=total_sku_margin,
    )
    clear_profit = _round_money(
        total_sku_margin - total_storage_cost - total_system_losses - tax_total
    )
    buyout = compute_buyout_coef_pct(sales_qty, returns_qty)
    return CfoEngineResult(
        kind=kind,
        tax_base_revenue=revenue_rrc,
        tax_total=tax_total,
        total_sku_margin=total_sku_margin,
        total_storage_cost=_round_money(total_storage_cost),
        total_system_losses=_round_money(total_system_losses),
        total_ad_spend=_round_money(total_ad_spend),
        credit_deductions=_round_money(credit_deductions),
        clear_profit=clear_profit,
        sales_qty=_round_money(sales_qty),
        returns_qty=_round_money(returns_qty),
        buyout_coef_pct=round(buyout, 1),
        sku_buckets=buckets,
        retail_price_source=retail_price_source,
    )


def _aggregate_cfo_transactions(
    matrix: list[list[str]],
    *,
    tax_preset: object | None = None,
    skip_detail_withholdings: bool = False,
    skip_detail_storage: bool = False,
) -> CfoEngineResult | None:
    cols = _resolve_cfo_tx_columns(matrix[0])
    if cols is None:
        return None

    buckets: dict[tuple[str, str], CfoSkuBucket] = {}
    tax_base_revenue = 0.0
    total_storage_cost = 0.0
    total_system_losses = 0.0
    total_ad_spend = 0.0
    credit_deductions = 0.0
    sales_qty = 0.0
    returns_qty = 0.0
    retail_sources: set[str] = set()

    for row in matrix[1:]:
        blob = _cfo_row_blob(row, cols)
        doc_type = _cfo_cell(row, cols.doc_type)
        kind = classify_cfo_tx_row(blob, doc_type=doc_type)
        if kind == "skip":
            continue
        if skip_detail_withholdings and _should_skip_detail_withholding_row(
            blob, doc_type, kind=kind
        ):
            continue
        if skip_detail_storage and _should_skip_detail_storage_row(blob, doc_type, kind=kind):
            continue

        retail_amount, src = _row_retail_amount(row, cols)
        if src != "missing":
            retail_sources.add(src)
        qty = abs(safe_float(row[cols.qty])) if cols.qty is not None and cols.qty < len(row) else 0.0
        loss_amount = _row_direct_amount(row, cols.payout_price) or abs(retail_amount)
        if kind in ("storage", "acceptance", "utilization"):
            total_storage_cost += loss_amount or _row_direct_amount(row, cols.commission)
            continue
        if kind == "ad":
            total_ad_spend += loss_amount
            continue
        if kind == "system_loss":
            total_system_losses += loss_amount
            if "кредит" in blob or "кредит" in doc_type.lower():
                credit_deductions += loss_amount
            continue

        name, article = _cfo_sku_identity(row, cols)
        from services.wb_transaction_parse import is_valid_wb_sku

        has_sku = is_valid_wb_sku(name, article)
        bucket = _cfo_get_bucket(buckets, name, article) if has_sku else None

        if kind in ("sale", "sale_adjustment"):
            signed = abs(retail_amount) if retail_amount else loss_amount
            unit_qty = qty if qty > 0 else (1.0 if signed > 0 else 0.0)
            if has_sku and bucket is not None:
                bucket.gross_sales_rrc += signed
                bucket.sales_qty += unit_qty
                bucket.deliveries_qty += unit_qty
                bucket.forward_logistics += _row_direct_amount(row, cols.forward_logistics)
                bucket.reverse_logistics += _row_direct_amount(row, cols.reverse_logistics)
                bucket.commission += _row_direct_amount(row, cols.commission)
                if cols.cost is not None and cols.cost < len(row):
                    bucket.cost_rub += abs(safe_float(row[cols.cost]))
                if cols.stock is not None and cols.stock < len(row):
                    bucket.stock_qty = max(bucket.stock_qty, max(0.0, safe_float(row[cols.stock])))
            tax_base_revenue += signed
            sales_qty += unit_qty
            continue

        if kind in ("return", "storno", "cancel"):
            signed = abs(retail_amount) if retail_amount else loss_amount
            unit_qty = qty if qty > 0 else (1.0 if signed > 0 else 0.0)
            if has_sku and bucket is not None:
                bucket.gross_sales_rrc -= signed
                bucket.returns_qty += unit_qty
                bucket.forward_logistics += _row_direct_amount(row, cols.forward_logistics)
                bucket.reverse_logistics += _row_direct_amount(row, cols.reverse_logistics)
                bucket.commission += _row_direct_amount(row, cols.commission)
            tax_base_revenue -= signed
            returns_qty += unit_qty
            continue

        if kind == "forward_logistics":
            amount = _row_direct_amount(row, cols.forward_logistics) or loss_amount
            if has_sku and bucket is not None:
                bucket.forward_logistics += amount
            continue
        if kind == "reverse_logistics":
            amount = _row_direct_amount(row, cols.reverse_logistics) or loss_amount
            if has_sku and bucket is not None:
                bucket.reverse_logistics += amount
            continue
        if kind == "commission":
            amount = _row_direct_amount(row, cols.commission) or loss_amount
            if has_sku and bucket is not None:
                bucket.commission += amount
            continue

    retail_source = "rrc" if "rrc" in retail_sources else (
        "payout_fallback" if "payout_fallback" in retail_sources else "missing"
    )
    return _finalize_cfo_engine(
        kind="transaction",
        buckets=buckets,
        tax_base_revenue=tax_base_revenue,
        total_storage_cost=total_storage_cost,
        total_system_losses=total_system_losses,
        total_ad_spend=total_ad_spend,
        credit_deductions=credit_deductions,
        sales_qty=sales_qty,
        returns_qty=returns_qty,
        retail_price_source=retail_source,
        tax_preset=tax_preset,
    )


def _aggregate_cfo_matrix(
    matrix: list[list[str]],
    *,
    platform: str | None,
    tax_preset: object | None = None,
) -> CfoEngineResult | None:
    from services.marketplace_platform import get_marketplace_profile
    from services.wb_transaction_parse import is_valid_wb_sku

    if not matrix or len(matrix) < 2:
        return None
    profile = get_marketplace_profile(platform)
    headers = matrix[0]

    name_col = _matrix_col(headers, ("предмет", "наименование", "номенклатур", "бренд", "товар")) or 0
    article_col = _matrix_col(headers, _WB_ARTICLE_HINTS)
    retail_col = find_column_index(headers, "retail_price")
    payout_col = find_column_index(headers, "payout_price")
    if retail_col is None:
        retail_col = _matrix_col(headers, profile.revenue_hints)
    sales_col = _matrix_col(headers, profile.sales_hints, require_qty=True) or _matrix_col(
        headers, profile.sales_hints
    )
    del_col = _matrix_col(headers, profile.delivery_hints, require_qty=True) or _matrix_col(
        headers, profile.delivery_hints
    )
    ret_col = _matrix_col(headers, profile.return_hints, require_qty=True) or _matrix_col(
        headers, profile.return_hints
    )
    comm_col = _matrix_col(headers, profile.commission_hints)
    return_log_cols = _matrix_return_logistics_cols(headers)
    forward_log_col = find_column_index(headers, "delivery", exclude_substrings=("хранен",))
    if forward_log_col is None:
        forward_log_col = _matrix_forward_logistics_col(headers, return_log_cols)
    cost_col = find_column_index(headers, "cost") or _matrix_col(headers, _MATRIX_COST_HINTS)
    stock_col = _matrix_col(headers, profile.stock_hints)
    ad_cols = [
        idx
        for idx, h in enumerate(headers)
        if any(x in (h or "").lower() for x in profile.ad_hints)
        and "кредит" not in (h or "").lower()
        and "хранен" not in (h or "").lower()
    ]

    buckets: dict[tuple[str, str], CfoSkuBucket] = {}
    tax_base_revenue = 0.0
    total_ad_spend = 0.0
    retail_sources: set[str] = set()

    for row in matrix[1:]:
        name, article = _row_sku_identity(row, name_col=name_col, article_col=article_col)
        if _is_total_row(name) or not is_valid_wb_sku(name, article):
            continue
        bucket = _cfo_get_bucket(buckets, name, article)

        retail_val = 0.0
        if retail_col is not None and retail_col < len(row):
            retail_val = safe_float(row[retail_col])
            if retail_val != 0:
                retail_sources.add("rrc")
        if retail_val == 0 and payout_col is not None and payout_col < len(row):
            retail_val = safe_float(row[payout_col])
            if retail_val != 0:
                retail_sources.add("payout_fallback")
        if retail_val > 0:
            bucket.gross_sales_rrc += retail_val
            tax_base_revenue += retail_val

        if sales_col is not None and sales_col < len(row):
            bucket.sales_qty += safe_float(row[sales_col])
        if del_col is not None and del_col < len(row):
            bucket.deliveries_qty += safe_float(row[del_col])
        if ret_col is not None and ret_col < len(row):
            bucket.returns_qty += safe_float(row[ret_col])
        if comm_col is not None and comm_col < len(row):
            bucket.commission += abs(safe_float(row[comm_col]))
        if forward_log_col is not None and forward_log_col < len(row):
            low_hdr = (headers[forward_log_col] or "").lower()
            if "хранен" not in low_hdr:
                bucket.forward_logistics += abs(safe_float(row[forward_log_col]))
        for rl_col in return_log_cols:
            if rl_col < len(row):
                bucket.reverse_logistics += abs(safe_float(row[rl_col]))
        if cost_col is not None and cost_col < len(row):
            bucket.cost_rub += abs(safe_float(row[cost_col]))
        if stock_col is not None and stock_col < len(row):
            bucket.stock_qty += max(0.0, safe_float(row[stock_col]))

    if not buckets:
        return None

    if total_ad_spend <= 0:
        for ac in ad_cols:
            for row in matrix[1:]:
                if ac < len(row):
                    total_ad_spend += abs(safe_float(row[ac]))

    sales_qty = sum(b.sales_qty for b in buckets.values())
    returns_qty = sum(b.returns_qty for b in buckets.values())
    if returns_qty > 0:
        deliveries_qty = sum(b.deliveries_qty for b in buckets.values())
        if deliveries_qty > 0:
            returns_qty = min(returns_qty, deliveries_qty)
        if sales_qty > 0:
            returns_qty = min(returns_qty, sales_qty * 2.0)

    retail_source = "rrc" if "rrc" in retail_sources else (
        "payout_fallback" if "payout_fallback" in retail_sources else "missing"
    )
    return _finalize_cfo_engine(
        kind="matrix",
        buckets=buckets,
        tax_base_revenue=tax_base_revenue,
        total_storage_cost=0.0,
        total_system_losses=0.0,
        total_ad_spend=total_ad_spend,
        credit_deductions=0.0,
        sales_qty=sales_qty,
        returns_qty=returns_qty,
        retail_price_source=retail_source,
        tax_preset=tax_preset,
    )


def aggregate_cfo_engine_v11_1(
    matrix: list[list[str]],
    *,
    platform: str | None = None,
    tax_preset_id: str | None = None,
    skip_detail_withholdings: bool = False,
    skip_detail_storage: bool = False,
) -> CfoEngineResult | None:
    """
    CFO Engine v11.1 — единая математическая модель агрегации WB/Ozon xlsx.

    Налоговая база — РРЦ; прямые расходы SKU; хранение и системные потери — глобально.
    """
    from services.audit_tax import resolve_audit_tax_preset

    if not matrix or len(matrix) < 2:
        return None
    preset = resolve_audit_tax_preset(tax_preset_id)
    tx = _aggregate_cfo_transactions(
        matrix,
        tax_preset=preset,
        skip_detail_withholdings=skip_detail_withholdings,
        skip_detail_storage=skip_detail_storage,
    )
    if tx is not None and (tx.sku_buckets or tx.total_storage_cost or tx.total_system_losses):
        return tx
    return _aggregate_cfo_matrix(matrix, platform=platform, tax_preset=preset)


_TOP_A_SHARE = 0.20
_OOS_RISK_DAYS = 5
_DEFAULT_REVERSE_LOGISTICS_RUB = 50.0
_RETURN_LOGISTICS_COL_HINTS: tuple[tuple[str, ...], ...] = (
    ("обратн", "логистик"),
    ("обратн", "перевоз"),
    ("обратн", "доставк"),
    ("логистик", "возврат"),
    ("перевоз", "возврат"),
    ("издерж", "возврат"),
)
_RETURN_LOGISTICS_PHRASES: tuple[str, ...] = (
    "обратная логистика",
    "логистика возврат",
    "логистика по возврат",
)
_WB_LABEL_HINTS = ("предмет", "артикул", "наименование", "номенклатур", "бренд")
_WB_NAME_HINTS = ("предмет", "наименование", "номенклатур", "бренд", "товар")
_WB_ARTICLE_HINTS = ("артикул", "sku", "nm id", "nmid", "vendor", "код товара", "barcode", "штрих")
_WB_REVENUE_HINTS = ("перечислению", "выруч", "заработок")
_WB_SALES_HINTS = ("выкупили", "реализован", "продаж")
_WB_DELIVERY_HINTS = ("доставк", "к клиенту")
_WB_RETURN_HINTS = ("возврат",)
_WB_RETURN_ID_SKIP = (
    "srid",
    "rrid",
    "rrd",
    "id ",
    " id",
    "номер",
    "код возврата",
    "документ",
    "транзак",
    "barcode",
    "штрих",
)


def _is_return_id_column(header: str) -> bool:
    """Колонки-ID транзакций не считаем количеством возвратов в штуках."""
    low = (header or "").lower()
    if any(q in low for q in _WB_QTY_UNIT_HINTS):
        return False
    return any(s in low for s in _WB_RETURN_ID_SKIP)
_WB_COMMISSION_HINTS = ("вознагражден", "комисс")
_WB_LOGISTICS_HINTS = ("логистик", "доставк", "хранен")
_WB_AD_HINTS = ("продвижен", "реклам", "удержан")
_WB_STOCK_HINTS = ("остаток", "склад", "stock", "quantity")
_WB_WAREHOUSE_HINTS = (
    "склад",
    "warehouse",
    "наименование склада",
    "офис",
    "логистический склад",
    "склад отгрузки",
)
_WB_VOLUME_HINTS = ("литраж", "объем", "объём", "volume", "литр", "габарит")
_MATRIX_COST_HINTS = ("себестоимость", "себестоим", "себес", "закупка", "закуп", "cost")


@dataclass(frozen=True)
class MatrixSkuDetail:
    """Товарная строка ETL: имя, артикул, выручка, маржа, выкуп."""

    name: str
    article_id: str
    revenue: float
    net_profit: float
    buyout_pct: float
    abc_group: str | None = None
    sales_qty: float = 0.0
    stock_qty: float = 0.0
    unit_cost_rub: float = 0.0

    @property
    def label(self) -> str:
        """Краткое имя для обратной совместимости."""
        return self.name

    def catalog_line(self) -> str:
        """Формат для промпта: Имя (Артикул) — Выручка — Маржа — Выкуп."""
        rev = f"{self.revenue:,.2f}".replace(",", " ")
        margin = f"{self.net_profit:,.2f}".replace(",", " ")
        return (
            f"{self.name} (Артикул: {self.article_id}) — "
            f"{rev} руб. — {margin} руб. — {self.buyout_pct:.1f}%"
        )


@dataclass(frozen=True)
class MatrixAbcSku:
    name: str
    article_id: str
    revenue: float
    net_profit: float
    buyout_pct: float
    abc_group: str

    @property
    def label(self) -> str:
        return self.name


@dataclass(frozen=True)
class MatrixOosForecast:
    label: str
    stock_qty: float
    sales_period_qty: float
    days_until_stockout: float | None
    risk_out_of_stock: bool


@dataclass(frozen=True)
class SellerMatrixEtl:
    """Результат локального ETL по строкам отчёта маркетплейса (0 ₽ OpenRouter)."""

    abc_group_a: tuple[MatrixAbcSku, ...]
    abc_group_c: tuple[MatrixAbcSku, ...]
    abc_a_leader: str
    logistics_fomo_rub: float
    logistics_fomo_detail: str
    oos_forecasts: tuple[MatrixOosForecast, ...]
    oos_critical_sku: str | None
    oos_critical_days: float | None
    logistics_fomo_items: tuple[str, ...] = ()
    reverse_logistics_shop_avg: float = 0.0
    return_logistics_block: str = ""
    sku_catalog: tuple[MatrixSkuDetail, ...] = ()
    outsider_sku: MatrixSkuDetail | None = None


def _matrix_col(headers: list[str], hints: tuple[str, ...], *, require_qty: bool = False) -> int | None:
    for idx, header in enumerate(headers):
        low = (header or "").lower()
        if not any(h in low for h in hints):
            continue
        if require_qty and not any(q in low for q in _WB_QTY_UNIT_HINTS):
            continue
        if any(h in low for h in _WB_RETURN_HINTS) and _is_return_id_column(header):
            continue
        return idx
    return None


def _matrix_warehouse_col(headers: list[str]) -> int | None:
    """Колонка названия склада WB (не путать с «остаток на складе»)."""
    for idx, header in enumerate(headers):
        low = (header or "").lower()
        if "остаток" in low:
            continue
        if any(h in low for h in _WB_WAREHOUSE_HINTS):
            return idx
    return None


def _is_total_row(label: str) -> bool:
    low = (label or "").strip().lower()
    return low.startswith(("итого", "всего", "total"))


def _matrix_name_and_article_cols(headers: list[str]) -> tuple[int, int | None]:
    """Колонка названия товара и опционально отдельный артикул."""
    article_col = _matrix_col(headers, _WB_ARTICLE_HINTS)
    name_col: int | None = None
    for hints in (_WB_NAME_HINTS, _WB_LABEL_HINTS):
        for idx, header in enumerate(headers):
            low = (header or "").lower()
            if any(h in low for h in hints) and idx != article_col:
                name_col = idx
                break
        if name_col is not None:
            break
    if name_col is None:
        name_col = _matrix_col(headers, _WB_LABEL_HINTS) or 0
    if article_col is not None and article_col == name_col:
        article_col = None
    return name_col, article_col


def resolve_wb_cfo_core_column_indices(
    headers: list[str],
) -> tuple[int | None, int | None]:
    """
    Критические колонки CFO v11.1: ``idx_sku`` (артикул) и ``idx_rrc`` (выручка/РРЦ).

    Колонки штрафов, кредитов и возвратов не обязательны — при отсутствии считаем 0.
    """
    tx_cols = _resolve_cfo_tx_columns(headers)
    if tx_cols is not None:
        idx_sku = tx_cols.article if tx_cols.article is not None else tx_cols.name
        idx_rrc = (
            tx_cols.retail_price
            if tx_cols.retail_price is not None
            else tx_cols.payout_price
        )
        return idx_sku, idx_rrc

    idx_sku = find_column_index(headers, "sku")
    if idx_sku is None:
        name_col, article_col = _matrix_name_and_article_cols(headers)
        idx_sku = article_col if article_col is not None else name_col

    idx_rrc = (
        find_column_index(headers, "retail_price")
        or find_column_index(headers, "payout_price")
        or find_column_index(headers, "sale_price")
        or _matrix_col(headers, _WB_REVENUE_HINTS)
    )
    return idx_sku, idx_rrc


def wb_core_finance_columns_recognized(matrix: list[list[str]]) -> bool:
    """Ключевые колонки SKU и выручки (РРЦ) распознаны по шапке отчёта."""
    if not matrix:
        return False
    headers = [str(h or "").strip() for h in matrix[0]]
    if not headers:
        return False
    idx_sku, idx_rrc = resolve_wb_cfo_core_column_indices(headers)
    return idx_sku is not None and idx_rrc is not None


def should_warn_column_structure(
    matrix: list[list[str]],
    *,
    revenue_total: float = 0.0,
) -> bool:
    """
    Предупреждение «не удалось распознать структуру колонок» — только если
    не найдены ``idx_sku`` или ``idx_rrc``. Штрафы/возвраты без колонок = 0.
    """
    _ = revenue_total  # совместимость вызовов
    if not matrix or len(matrix) < 2:
        return True
    headers = [str(h or "").strip() for h in matrix[0]]
    idx_sku, idx_rrc = resolve_wb_cfo_core_column_indices(headers)
    return idx_sku is None or idx_rrc is None


def _row_sku_identity(
    row: list[str],
    *,
    name_col: int,
    article_col: int | None,
) -> tuple[str, str]:
    name = (row[name_col] if name_col < len(row) else "").strip() or "—"
    if article_col is not None and article_col < len(row):
        article = (row[article_col] or "").strip() or name
    else:
        article = name
    if _is_technical_sku_cell(name) and _is_technical_sku_cell(article):
        return "—", "—"
    if _is_technical_sku_cell(name) and not _is_technical_sku_cell(article):
        name = article
    elif _is_technical_sku_cell(article) and not _is_technical_sku_cell(name):
        article = name
    return name[:64], article[:48]


@dataclass
class _SkuBucket:
    name: str
    article_id: str
    revenue: float = 0.0
    commission: float = 0.0
    logistics: float = 0.0
    ad_cost: float = 0.0
    extra_cost: float = 0.0
    cost_rub: float = 0.0
    sales_qty: float = 0.0
    deliveries_qty: float = 0.0
    returns_qty: float = 0.0
    stock_qty: float = 0.0
    return_logistics_rub: float = 0.0
    warehouse_name: str = ""
    volume_liters: float = 0.0

    @property
    def unit_cost_rub(self) -> float:
        if self.cost_rub <= 0:
            return 0.0
        if self.sales_qty > 0:
            return self.cost_rub / self.sales_qty
        if self.stock_qty > 0:
            return self.cost_rub / self.stock_qty
        return self.cost_rub

    @property
    def net_profit(self) -> float:
        """Маржа SKU = продажи − себестоимость − прямые расходы (комиссия + логистика)."""
        return (
            self.revenue
            - self.commission
            - self.logistics
            - self.return_logistics_rub
            - self.cost_rub
        )

    @property
    def unit_logistics(self) -> float:
        if self.sales_qty > 0:
            return self.logistics / self.sales_qty
        if self.deliveries_qty > 0:
            return self.logistics / self.deliveries_qty
        return 0.0

    @property
    def buyout_pct(self) -> float:
        return compute_buyout_coef_pct(self.sales_qty, self.returns_qty)

    def to_detail(self, *, abc_group: str | None = None) -> MatrixSkuDetail:
        return MatrixSkuDetail(
            name=self.name,
            article_id=self.article_id,
            revenue=round(self.revenue, 2),
            net_profit=round(self.net_profit, 2),
            buyout_pct=round(self.buyout_pct, 1),
            abc_group=abc_group,
            sales_qty=round(self.sales_qty, 2),
            stock_qty=round(self.stock_qty, 2),
            unit_cost_rub=round(self.unit_cost_rub, 2),
        )


def _matrix_return_logistics_cols(headers: list[str]) -> list[int]:
    """Колонки затрат именно на обратную логистику / возвраты (WB xlsx)."""
    cols: list[int] = []
    for idx, header in enumerate(headers):
        low = (header or "").lower()
        if any(phrase in low for phrase in _RETURN_LOGISTICS_PHRASES):
            cols.append(idx)
            continue
        if all(any(h in low for h in group) for group in _RETURN_LOGISTICS_COL_HINTS):
            cols.append(idx)
    return cols


def _matrix_forward_logistics_col(
    headers: list[str],
    return_log_cols: list[int],
) -> int | None:
    """Общая колонка логистики без признаков возврата."""
    for idx, header in enumerate(headers):
        if idx in return_log_cols:
            continue
        low = (header or "").lower()
        if any(h in low for h in ("логистик", "доставк", "хранен", "перевоз")):
            if "возврат" not in low and "обратн" not in low:
                return idx
    return _matrix_col(headers, _WB_LOGISTICS_HINTS)


def _validated_sku_returns_qty(bucket: _SkuBucket) -> float:
    """Фактические возвраты в штуках: не больше доставок/заказов по SKU."""
    raw = max(0.0, bucket.returns_qty)
    if raw <= 0:
        return 0.0
    upper = bucket.deliveries_qty if bucket.deliveries_qty > 0 else 0.0
    if bucket.sales_qty > 0:
        if upper <= 0:
            upper = bucket.sales_qty
        raw = min(raw, upper)
    elif upper > 0:
        raw = min(raw, upper)
    return raw


def _effective_reverse_logistics_qty(bucket: _SkuBucket) -> float:
    """
    Штуки для расчёта обратной логистики: приоритет — колонка «Возвраты, шт.»,
    иначе невыкупы (доставки − выкупы) с тем же потолком по доставкам.
    """
    returns = _validated_sku_returns_qty(bucket)
    if returns > 0:
        return returns
    if bucket.deliveries_qty <= 0:
        return 0.0
    non_buyout = max(0.0, bucket.deliveries_qty - bucket.sales_qty)
    return min(non_buyout, bucket.deliveries_qty)


def _non_buyout_qty(bucket: _SkuBucket) -> float:
    """Обратная совместимость: см. :func:`_effective_reverse_logistics_qty`."""
    return _effective_reverse_logistics_qty(bucket)


def _reverse_logistics_unit_rub(bucket: _SkuBucket, qty: float) -> float:
    """
    Средняя стоимость одной обратной логистики для SKU (руб./ед.).

    Приоритет: факт из отчёта → ночной кэш WB по складу и литражу → доля общей логистики.
    Живые запросы к API при разборе Excel **запрещены**.
    """
    if qty <= 0:
        return 0.0
    if bucket.return_logistics_rub > 0:
        return max(_DEFAULT_REVERSE_LOGISTICS_RUB, bucket.return_logistics_rub / qty)

    from services.wb_tariffs_cache import estimate_return_logistics_unit_rub

    cached_unit = estimate_return_logistics_unit_rub(
        bucket.warehouse_name,
        bucket.volume_liters or 1.0,
        floor_rub=_DEFAULT_REVERSE_LOGISTICS_RUB,
    )
    if bucket.warehouse_name.strip():
        return cached_unit

    if bucket.logistics > 0:
        if bucket.deliveries_qty > 0:
            return_share = min(1.0, qty / bucket.deliveries_qty)
            unit = (bucket.logistics * return_share) / qty
        else:
            unit = bucket.logistics / qty
        return max(_DEFAULT_REVERSE_LOGISTICS_RUB, unit)
    if cached_unit > _DEFAULT_REVERSE_LOGISTICS_RUB:
        return cached_unit
    return _DEFAULT_REVERSE_LOGISTICS_RUB


def _format_return_logistics_fomo_line(
    name: str,
    article_id: str,
    *,
    returns_count: float,
    total_loss_rub: float,
    warehouse_name: str = "",
    volume_liters: float = 0.0,
    unit_rub: float = 0.0,
) -> str:
    label = _format_sku_label_for_json(name, article_id)
    loss_s = f"{total_loss_rub:,.2f}".replace(",", " ")
    wh = (warehouse_name or "").strip()
    vol = float(volume_liters or 0.0)
    if wh and vol > 0 and unit_rub > 0:
        return (
            f"Логистика возвратов: {label} ({wh}, {vol:.1f} л × "
            f"{unit_rub:,.2f} руб./шт.): {returns_count:.0f} возвратов. "
            f"Убыток на покатушках (кэш тарифов WB): ≈ {loss_s} руб."
        )
    return (
        f"Логистика возвратов: {label}: {returns_count:.0f} возвратов. "
        f"Общий убыток на пустых покатушках: ≈ {loss_s} руб."
    )


def compute_seller_matrix_etl(
    rows: list[list[str]],
    *,
    revenue_total: float = 0.0,
    platform: str | None = None,
) -> SellerMatrixEtl | None:
    """
    ABC по чистой прибыли SKU, FOMO логистики невыкупов, прогноз OOS.

    ``rows`` — матрица ``[headers, *data]`` из xlsx/csv (см. :func:`read_xlsx_rows_from_path`).
    ``platform`` — wildberries | ozon | yandex | 1c (формула P&L площадки).
    """
    from services.marketplace_platform import get_marketplace_profile, normalize_marketplace_platform

    if not rows or len(rows) < 2:
        return None

    profile = get_marketplace_profile(platform)
    platform_id = normalize_marketplace_platform(platform)

    headers = [str(h).strip() for h in rows[0]]
    cost_col: int | None = find_column_index(headers, "cost")
    name_col, article_col = _matrix_name_and_article_cols(headers)
    sku_col = find_column_index(headers, "sku")
    if sku_col is not None and sku_col != name_col:
        article_col = sku_col
    rev_col = find_column_index(headers, "sale_price") or _matrix_col(headers, profile.revenue_hints)
    sales_col = _matrix_col(headers, profile.sales_hints, require_qty=True)
    if sales_col is None:
        sales_col = _matrix_col(headers, profile.sales_hints)
    del_col = _matrix_col(headers, profile.delivery_hints, require_qty=True)
    if del_col is None:
        del_col = _matrix_col(headers, profile.delivery_hints)
    ret_col = _matrix_col(headers, profile.return_hints, require_qty=True)
    if ret_col is None:
        ret_col = _matrix_col(headers, profile.return_hints)
    comm_col = _matrix_col(headers, profile.commission_hints)
    return_log_cols = _matrix_return_logistics_cols(headers)
    sem_delivery = find_column_index(headers, "delivery", exclude_substrings=("хранен",))
    log_col = sem_delivery if sem_delivery is not None else _matrix_forward_logistics_col(
        headers, return_log_cols
    )
    if log_col is None:
        log_col = _matrix_col(headers, profile.logistics_hints)
    if cost_col is None:
        cost_col = _matrix_col(headers, _MATRIX_COST_HINTS)
    if cost_col is None:
        logger.warning(
            "compute_seller_matrix_etl: колонка себестоимости не найдена "
            "(подсказки %s); себестоимость принимается за 0",
            WB_COLUMN_SYNONYMS.get("cost", _MATRIX_COST_HINTS),
        )
    ad_cols = [
        idx
        for idx, h in enumerate(headers)
        if any(x in (h or "").lower() for x in profile.ad_hints)
    ]
    extra_cols = [
        idx
        for idx, h in enumerate(headers)
        if any(x in (h or "").lower() for x in profile.extra_deduction_hints)
        and idx not in ad_cols
        and idx != comm_col
        and idx != log_col
        and (cost_col is not None and idx != cost_col)
    ]
    stock_col = _matrix_col(headers, profile.stock_hints)
    warehouse_col = find_column_index(headers, "warehouse")
    if warehouse_col is None:
        warehouse_col = _matrix_warehouse_col(headers)
    volume_col = find_column_index(headers, "volume_liters") or _matrix_col(
        headers, _WB_VOLUME_HINTS
    )

    from services.wb_report_parser import parse_wb_report
    from services.wb_transaction_parse import is_valid_wb_sku

    report = parse_wb_report(rows, platform=platform)
    buckets: dict[tuple[str, str], _SkuBucket] = {}
    if report is not None:
        for (name, article), sm in report.sku_by_key.items():
            buckets[(name, article)] = _SkuBucket(
                name=sm.name,
                article_id=sm.article_id,
                revenue=sm.revenue,
                sales_qty=sm.sales_qty,
                deliveries_qty=sm.deliveries_qty,
                returns_qty=sm.returns_qty,
                logistics=sm.logistics,
                commission=sm.commission,
                cost_rub=sm.cost_rub,
                stock_qty=sm.stock_qty,
                return_logistics_rub=sm.return_logistics_rub,
            )
    else:
        for row in rows[1:]:
            name, article_id = _row_sku_identity(row, name_col=name_col, article_col=article_col)
            if _is_total_row(name):
                continue
            if not is_valid_wb_sku(name, article_id):
                continue
            bucket = buckets.get((name, article_id))
            if bucket is None:
                bucket = _SkuBucket(name=name, article_id=article_id)
                buckets[(name, article_id)] = bucket
            if rev_col is not None and rev_col < len(row):
                val = safe_float(row[rev_col])
                if val > 0:
                    bucket.revenue += val
            if sales_col is not None and sales_col < len(row):
                bucket.sales_qty += safe_float(row[sales_col])
            if del_col is not None and del_col < len(row):
                bucket.deliveries_qty += safe_float(row[del_col])
            if ret_col is not None and ret_col < len(row):
                bucket.returns_qty += safe_float(row[ret_col])
            if comm_col is not None and comm_col < len(row):
                bucket.commission += abs(safe_float(row[comm_col]))
            if log_col is not None and log_col < len(row):
                low_hdr = (headers[log_col] or "").lower()
                if "хранен" not in low_hdr:
                    bucket.logistics += abs(safe_float(row[log_col]))
            for rl_col in return_log_cols:
                if rl_col < len(row):
                    bucket.return_logistics_rub += abs(safe_float(row[rl_col]))
            for ac in ad_cols:
                low_hdr = (headers[ac] or "").lower()
                if "кредит" in low_hdr or "хранен" in low_hdr:
                    continue
                if ac < len(row):
                    bucket.ad_cost += abs(safe_float(row[ac]))
            for ec in extra_cols:
                if ec < len(row):
                    bucket.extra_cost += abs(safe_float(row[ec]))
            if cost_col is not None and cost_col < len(row):
                bucket.cost_rub += abs(safe_float(row[cost_col]))
            if stock_col is not None and stock_col < len(row):
                bucket.stock_qty += max(0.0, safe_float(row[stock_col]))
            if warehouse_col is not None and warehouse_col < len(row):
                wh = str(row[warehouse_col] or "").strip()
                if wh and not bucket.warehouse_name:
                    bucket.warehouse_name = wh
            if volume_col is not None and volume_col < len(row):
                vol = safe_float(row[volume_col])
                if vol > 0:
                    bucket.volume_liters = max(bucket.volume_liters, vol)

    if report is not None and (warehouse_col is not None or volume_col is not None):
        for row in rows[1:]:
            name, article_id = _row_sku_identity(row, name_col=name_col, article_col=article_col)
            if _is_total_row(name):
                continue
            bucket = buckets.get((name, article_id))
            if bucket is None:
                continue
            if warehouse_col is not None and warehouse_col < len(row):
                wh = str(row[warehouse_col] or "").strip()
                if wh and not bucket.warehouse_name:
                    bucket.warehouse_name = wh
            if volume_col is not None and volume_col < len(row):
                vol = safe_float(row[volume_col])
                if vol > 0:
                    bucket.volume_liters = max(bucket.volume_liters, vol)

    buckets = {k: v for k, v in buckets.items() if is_valid_wb_sku(k[0], k[1])}

    if not buckets:
        return None

    # ABC по чистой прибыли (Парето)
    ranked = sorted(buckets.items(), key=lambda x: x[1].net_profit, reverse=True)
    n = len(ranked)
    top_a_n = max(1, math.ceil(n * _TOP_A_SHARE))
    group_a_keys = {key for key, _ in ranked[:top_a_n]}
    group_a = tuple(
        MatrixAbcSku(
            name=b.name,
            article_id=b.article_id,
            revenue=round(b.revenue, 2),
            net_profit=round(b.net_profit, 2),
            buyout_pct=round(b.buyout_pct, 1),
            abc_group="A",
        )
        for key, b in ranked
        if key in group_a_keys
    )
    group_c = tuple(
        MatrixAbcSku(
            name=b.name,
            article_id=b.article_id,
            revenue=round(b.revenue, 2),
            net_profit=round(b.net_profit, 2),
            buyout_pct=round(b.buyout_pct, 1),
            abc_group="C",
        )
        for key, b in buckets.items()
        if b.net_profit <= 0
    )
    abc_leader = group_a[0].name if group_a else (ranked[0][1].name if ranked else "—")

    sku_catalog = tuple(
        b.to_detail(
            abc_group=(
                "A"
                if key in group_a_keys
                else ("C" if b.net_profit <= 0 else "B")
            )
        )
        for key, b in ranked
    )
    outsider_sku: MatrixSkuDetail | None = None
    if group_c:
        worst_c = min(group_c, key=lambda s: s.net_profit)
        outsider_sku = MatrixSkuDetail(
            name=worst_c.name,
            article_id=worst_c.article_id,
            revenue=worst_c.revenue,
            net_profit=worst_c.net_profit,
            buyout_pct=worst_c.buyout_pct,
            abc_group="C",
        )

    # FOMO: обратная логистика невыкупов — реальная средняя стоимость из xlsx
    logistics_fomo = 0.0
    fomo_parts: list[str] = []
    weighted_unit_sum = 0.0
    weighted_unit_qty = 0.0
    for (name, _), b in buckets.items():
        returns_qty = _effective_reverse_logistics_qty(b)
        if returns_qty <= 0:
            continue
        unit_log = _reverse_logistics_unit_rub(b, returns_qty)
        loss = returns_qty * unit_log
        if loss > 0:
            logistics_fomo += loss
            weighted_unit_sum += unit_log * returns_qty
            weighted_unit_qty += returns_qty
            if len(fomo_parts) < 6:
                fomo_parts.append(
                    _format_return_logistics_fomo_line(
                        name,
                        b.article_id,
                        returns_count=returns_qty,
                        total_loss_rub=loss,
                        warehouse_name=b.warehouse_name,
                        volume_liters=b.volume_liters,
                        unit_rub=unit_log,
                    )
                )
    reverse_logistics_shop_avg = (
        weighted_unit_sum / weighted_unit_qty if weighted_unit_qty > 0 else 0.0
    )
    return_logistics_block = (
        "\n".join(f"• {line}" for line in fomo_parts)
        if fomo_parts
        else "• существенных потерь на обратной логистике не выявлено"
    )
    logistics_detail = (
        "Логистика возвратов: " + "; ".join(fomo_parts)
        if fomo_parts
        else "Существенных потерь на обратной логистике не выявлено."
    )

    # OOS: остаток / (продажи за период / 7 дней)
    oos_list: list[MatrixOosForecast] = []
    for (name, _), b in buckets.items():
        daily = b.sales_qty / 7.0 if b.sales_qty > 0 else 0.0
        days: float | None = None
        risk = False
        if daily > 0 and b.stock_qty > 0:
            days = b.stock_qty / daily
            risk = days < _OOS_RISK_DAYS
        elif b.stock_qty <= 0 and b.sales_qty > 0:
            days = 0.0
            risk = True
        oos_list.append(
            MatrixOosForecast(
                label=name,
                stock_qty=b.stock_qty,
                sales_period_qty=b.sales_qty,
                days_until_stockout=round(days, 1) if days is not None else None,
                risk_out_of_stock=risk,
            )
        )
    risky = [f for f in oos_list if f.risk_out_of_stock and f.days_until_stockout is not None]
    risky.sort(key=lambda x: x.days_until_stockout or 999.0)
    oos_sku: str | None = None
    oos_days: float | None = None
    if risky:
        oos_sku = risky[0].label
        oos_days = risky[0].days_until_stockout

    _ = revenue_total  # зарезервировано для будущей нормализации долей
    return SellerMatrixEtl(
        abc_group_a=group_a,
        abc_group_c=group_c,
        abc_a_leader=abc_leader,
        logistics_fomo_rub=round(logistics_fomo, 2),
        logistics_fomo_detail=logistics_detail,
        logistics_fomo_items=tuple(fomo_parts),
        reverse_logistics_shop_avg=round(reverse_logistics_shop_avg, 2),
        return_logistics_block=return_logistics_block,
        oos_forecasts=tuple(oos_list),
        oos_critical_sku=oos_sku,
        oos_critical_days=oos_days,
        sku_catalog=sku_catalog,
        outsider_sku=outsider_sku,
    )


def _format_sku_label_for_json(name: str, article_id: str) -> str:
    name = (name or "").strip() or "—"
    article_id = (article_id or "").strip()
    if article_id and article_id != name and article_id != "—":
        return f"{name} (арт. {article_id})"
    return name


def _load_cfo_matrix_from_path(file_path: str) -> list[list[str]]:
    """Читает .xlsx / .csv с диска в матрицу строк (потоково для Excel)."""
    return load_cfo_workbook_from_path(file_path).matrix


def _cfo_engine_to_sku_data(engine: CfoEngineResult) -> dict[str, dict[str, float | int]]:
    """Словарь поартикульных метрик (формат интеграций CFO)."""
    out: dict[str, dict[str, float | int]] = {}
    for _key, bucket in engine.sku_buckets.items():
        sku = (bucket.article_id or bucket.name or "").strip()
        if not sku or sku == "—" or len(sku) < 2:
            continue
        payout_est = bucket.gross_sales_rrc - bucket.commission
        entry: dict[str, float | int | str] = {
            "sales_count": int(bucket.sales_qty),
            "returns_count": int(bucket.returns_qty),
            "rrc_revenue": _round_money(bucket.gross_sales_rrc),
            "payout": _round_money(payout_est),
            "delivery": _round_money(bucket.logistics),
            "stock": int(bucket.stock_qty),
        }
        human_name = (bucket.name or "").strip()
        if human_name and human_name not in ("—", "-", "–"):
            entry["human_name"] = human_name
        sales_qty = float(bucket.sales_qty)
        entry["net_profit_rub"] = _round_money(bucket.margin_rub)
        entry["unit_profit_rub"] = (
            _round_money(bucket.margin_rub / sales_qty) if sales_qty > 0 else 0.0
        )
        out[sku] = entry
    return out


def _cfo_oos_lists_from_etl(
    etl: SellerMatrixEtl | None,
    *,
    critical_days: int = _OOS_CRITICAL_DAYS,
) -> tuple[list[dict[str, object]], list[dict[str, object]]]:
    """
    OOS cfo-v12: цветовые статусы без штук и дней в тексте.

    🔴 остаток 0 (в т.ч. без продаж в периоде) · 🟡 критический запас < N дн.
    """
    if etl is None:
        return [], []

    catalog_map = {s.name: s for s in etl.sku_catalog}
    zero_stock: list[dict[str, object]] = []
    critical: list[dict[str, object]] = []

    for forecast in etl.oos_forecasts:
        sales = float(forecast.sales_period_qty)
        stock = float(forecast.stock_qty)
        detail = catalog_map.get(forecast.label)
        name = (detail.name if detail else forecast.label) or forecast.label
        article_id = (detail.article_id if detail else forecast.label) or forecast.label
        label = _format_sku_label_for_json(name, article_id)
        base = {
            "name": name,
            "article_id": article_id,
            "label": label,
            "sku": article_id,
        }
        if stock <= 0:
            zero_stock.append({**base, "stock_qty": 0, "sales_qty": round(sales, 2)})
            continue
        days = forecast.days_until_stockout
        if (
            sales > 0
            and forecast.risk_out_of_stock
            and days is not None
            and float(days) <= critical_days
        ):
            critical.append(
                {
                    **base,
                    "stock_qty": round(stock, 2),
                    "days_until_stockout": round(float(days), 1),
                }
            )
    critical.sort(
        key=lambda item: float(item.get("days_until_stockout") or 999.0),
    )
    return zero_stock, critical


def build_report_metrics_for_history(
    rows: list[list[str]],
    *,
    revenue_total: float = 0.0,
    platform: str | None = "wildberries",
    tax_type: str = "USN",
    tax_rate: float = 6.0,
    tax_preset_id: str | None = None,
) -> dict[str, object]:
    """
    Словарь для ``save_user_report_to_db``: CFO-метрики + ``final_metrics_json``.
    """
    from services.audit_tax import preset_from_regime_rate, resolve_audit_tax_preset

    if not rows or len(rows) < 2:
        return {"error": "empty_file", "cfo_build": _CFO_BUILD}

    preset = (
        resolve_audit_tax_preset(tax_preset_id)
        if tax_preset_id
        else preset_from_regime_rate(tax_type, tax_rate)
    )
    cfo = build_cfo_metrics_dict_from_rows(
        rows,
        platform or "wildberries",
        preset.regime,
        preset.rate_percent,
    )
    if cfo.get("error"):
        return cfo

    rev = float(revenue_total or cfo.get("total_revenue") or 0.0)
    final = build_final_metrics_json(
        rows,
        revenue_total=rev,
        platform=platform,
        tax_preset_id=preset.id,
    )
    pack: dict[str, object] = dict(cfo)
    pack["platform"] = platform or "wildberries"
    pack["final_metrics_json"] = final
    pack["cfo_build"] = _CFO_BUILD
    return pack


def collect_supply_chain_audit_from_rows(
    matrix: list[list[str]],
) -> dict[str, object]:
    """
    Операционный аудит поставок: топ регионов/складов и отмены по артикулам.

    Использует ``collections.Counter`` по строкам детализации WB.
    """
    empty: dict[str, object] = {
        "top_regions": [],
        "top_warehouses": [],
        "canceled_skus": [],
    }
    if not matrix or len(matrix) < 2:
        return empty

    headers = [str(h) for h in matrix[0]]
    warehouse_col = find_column_index(headers, "warehouse")
    if warehouse_col is None:
        warehouse_col = _matrix_warehouse_col(headers)
    region_col = find_column_index(headers, "region")
    sku_col = find_column_index(headers, "sku")
    op_col = find_column_index(headers, "op_type")
    name_col, article_col = _matrix_name_and_article_cols(headers)
    tx_cols = _resolve_cfo_tx_columns(headers)

    region_counter: Counter[str] = Counter()
    warehouse_counter: Counter[str] = Counter()
    canceled_skus: list[str] = []
    canceled_seen: set[str] = set()

    _OUTBOUND_KINDS = frozenset(
        {"sale", "sale_adjustment", "forward_logistics", "commission", "other"}
    )
    _SKIP_SUPPLY_KINDS = frozenset(
        {
            "skip",
            "storage",
            "acceptance",
            "utilization",
            "ad",
            "system_loss",
            "return",
            "storno",
            "cancel",
        }
    )

    for row in matrix[1:]:
        if not any(str(cell or "").strip() for cell in row):
            continue

        op_text = ""
        blob = ""
        if tx_cols is not None:
            blob = _cfo_row_blob(row, tx_cols)
            op_text = _cfo_cell(row, tx_cols.doc_type)
            kind = classify_cfo_tx_row(blob, doc_type=op_text)
        elif op_col is not None and op_col < len(row):
            op_text = str(row[op_col] or "").strip()
            blob = op_text.lower()
            kind = classify_cfo_tx_row(blob, doc_type=op_text)
        else:
            kind = "other"

        sku_label = ""
        if tx_cols is not None:
            name, article = _cfo_sku_identity(row, tx_cols)
            sku_label = (article or name or "").strip()
        else:
            if sku_col is not None and sku_col < len(row):
                sku_label = str(row[sku_col] or "").strip()
            elif article_col is not None and article_col < len(row):
                sku_label = str(row[article_col] or "").strip()
            elif name_col is not None and name_col < len(row):
                sku_label = str(row[name_col] or "").strip()

        op_low = f"{op_text} {blob}".lower()
        if "возврат" in op_low or "сторно" in op_low:
            if sku_label and sku_label not in canceled_seen:
                canceled_seen.add(sku_label)
                canceled_skus.append(sku_label)

        if kind in _SKIP_SUPPLY_KINDS:
            continue

        count_supply = kind in _OUTBOUND_KINDS
        if not count_supply:
            continue

        if warehouse_col is not None and warehouse_col < len(row):
            warehouse = map_wb_warehouse_label(str(row[warehouse_col] or "").strip())
            if warehouse:
                warehouse_counter[warehouse] += 1

        if region_col is not None and region_col < len(row):
            region = str(row[region_col] or "").strip()
            if region:
                region_counter[region] += 1

    return {
        "top_regions": [name for name, _ in region_counter.most_common(3)],
        "top_warehouses": [name for name, _ in warehouse_counter.most_common(2)],
        "canceled_skus": canceled_skus,
    }


def build_cfo_metrics_dict_from_rows(
    rows: list[list[str]],
    audit_platform: str,
    tax_type: str,
    tax_rate: float,
    *,
    aux_storage_cost: float = 0.0,
    aux_system_losses: float = 0.0,
    aux_storage_from_sheet: bool = False,
    aux_system_from_sheet: bool = False,
) -> dict[str, object]:
    """CFO Engine v11.1: матрица строк → метрики с динамическим налогом."""
    from services.audit_tax import preset_from_regime_rate

    if not rows or len(rows) < 2:
        return {
            "error": "empty_file",
            "cfo_build": _CFO_BUILD,
            "tax_type": tax_type,
            "tax_rate": tax_rate,
        }

    if not validate_wb_finance_detail_structure(rows):
        return wb_finance_invalid_structure_payload(
            tax_type=tax_type,
            tax_rate=tax_rate,
        )

    preset = preset_from_regime_rate(tax_type, tax_rate)
    engine = aggregate_cfo_engine_v11_1(
        rows,
        platform=audit_platform,
        tax_preset_id=preset.id,
        skip_detail_withholdings=aux_system_from_sheet,
        skip_detail_storage=aux_storage_from_sheet,
    )
    if engine is None:
        return {
            "error": "unparsed_report",
            "cfo_build": _CFO_BUILD,
            "tax_type": tax_type,
            "tax_rate": tax_rate,
        }

    if aux_storage_cost == 0.0 and aux_system_losses == 0.0:
        scanned_storage, scanned_system = scan_matrix_deep_costs(
            rows,
            skip_withholding_rows=aux_system_from_sheet,
            skip_storage_rows=aux_storage_from_sheet,
        )
        aux_storage_cost = scanned_storage
        aux_system_losses = scanned_system

    engine = apply_deep_workbook_costs_to_engine(
        engine,
        aux_storage_cost=aux_storage_cost,
        aux_system_losses=aux_system_losses,
        aux_storage_from_sheet=aux_storage_from_sheet,
        aux_system_from_sheet=aux_system_from_sheet,
    )
    revenue = engine.tax_base_revenue
    etl = compute_seller_matrix_etl(rows, revenue_total=revenue, platform=audit_platform)
    sku_data = _cfo_engine_to_sku_data(engine)
    oos_zero, oos_critical = _cfo_oos_lists_from_etl(etl)
    margin_pct = (
        round(engine.clear_profit / revenue * 100.0, 2) if revenue > 0 else 0.0
    )
    supply_audit = collect_supply_chain_audit_from_rows(rows)

    return {
        "cfo_build": _CFO_BUILD,
        "engine": CFO_ENGINE_NAME,
        "tax_type": preset.regime,
        "tax_rate": preset.rate_percent,
        "total_revenue": revenue,
        "tax_total": engine.tax_total,
        "net_profit": engine.clear_profit,
        "margin_pct": margin_pct,
        "sku_data": sku_data,
        "total_storage_cost": engine.total_storage_cost,
        "total_system_losses": engine.total_system_losses,
        "total_sku_margin": engine.total_sku_margin,
        "oos_zero_stock_items": oos_zero,
        "oos_critical_sku": oos_critical,
        "top_regions": supply_audit["top_regions"],
        "top_warehouses": supply_audit["top_warehouses"],
        "canceled_skus": supply_audit["canceled_skus"],
    }


def sync_table_cfo_processing_worker(
    file_path: str,
    audit_platform: str,
    table_subrole: str,
    tax_type: str,
    tax_rate: float,
) -> dict[str, object]:
    """
    Высокоскоростной бухгалтерский ETL-парсер ядра CFO Engine v11.1.

    Читает xlsx/csv с диска, агрегирует транзакции WB и возвращает метрики
    с динамическим налогом (``tax_type`` / ``tax_rate`` из FSM аудита).
    """
    _ = table_subrole  # зарезервировано для ветвления подролей table_generator
    loaded = load_cfo_workbook_from_path(file_path)
    return build_cfo_metrics_dict_from_rows(
        loaded.matrix,
        audit_platform,
        tax_type,
        tax_rate,
        aux_storage_cost=loaded.aux_storage_cost,
        aux_system_losses=loaded.aux_system_losses,
        aux_storage_from_sheet=loaded.storage_from_dedicated_sheet,
        aux_system_from_sheet=loaded.system_from_dedicated_sheet,
    )


def build_final_metrics_json(
    rows: list[list[str]],
    *,
    revenue_total: float,
    platform: str | None = None,
    tax_preset_id: str | None = None,
) -> dict[str, Any]:
    """
    CFO Engine v11.1: все математические метрики в Python.

    OpenRouter получает только этот JSON — без пересчёта на стороне LLM.
    """
    if revenue_total <= 0 or not rows or len(rows) < 2:
        return {"error": "empty_or_no_revenue", "cfo_build": _CFO_BUILD, "engine": CFO_ENGINE_NAME}

    engine = aggregate_cfo_engine_v11_1(
        rows, platform=platform, tax_preset_id=tax_preset_id
    )
    etl = compute_seller_matrix_etl(rows, revenue_total=revenue_total, platform=platform)

    if engine is None:
        return {"error": "unparsed_report", "cfo_build": _CFO_BUILD, "engine": CFO_ENGINE_NAME}

    tax_base = engine.tax_base_revenue if engine.tax_base_revenue > 0 else revenue_total
    tax_usn = engine.tax_total if engine.tax_total > 0 else _round_money(tax_base * _USN_RATE)
    clear_profit = engine.clear_profit
    operational_profit = engine.operational_profit
    margin_rate = (
        round(clear_profit / tax_base * 100.0, 1) if tax_base > 0 else 0.0
    )
    drr_pct = (
        round(engine.total_ad_spend / tax_base * 100.0, 1)
        if tax_base > 0 and engine.total_ad_spend > 0
        else 0.0
    )

    sku_rows: list[dict[str, Any]] = []
    group_a: list[str] = []
    group_c: list[str] = []
    if etl:
        group_a = [
            _format_sku_label_for_json(s.name, s.article_id) for s in etl.abc_group_a
        ]
        group_c = [
            _format_sku_label_for_json(s.name, s.article_id) for s in etl.abc_group_c
        ]
        for detail in etl.sku_catalog:
            sku_rows.append(
                {
                    "name": detail.name,
                    "article_id": detail.article_id,
                    "label": _format_sku_label_for_json(detail.name, detail.article_id),
                    "revenue_rub": round(detail.revenue, 2),
                    "net_profit_rub": round(detail.net_profit, 2),
                    "margin_rub": round(detail.net_profit, 2),
                    "buyout_pct": round(detail.buyout_pct, 1),
                    "sales_qty": round(detail.sales_qty, 2),
                    "stock_qty": round(detail.stock_qty, 2),
                    "unit_cost_rub": round(detail.unit_cost_rub, 2),
                    "abc_group": detail.abc_group,
                }
            )

    oos_predictions: dict[str, int] = {}
    if etl:
        for forecast in etl.oos_forecasts:
            if forecast.risk_out_of_stock and forecast.days_until_stockout is not None:
                oos_predictions[forecast.label] = max(0, int(forecast.days_until_stockout))

    return {
        "cfo_build": _CFO_BUILD,
        "engine": CFO_ENGINE_NAME,
        "parser": "wb_final_metrics_v11_1",
        "retail_price_source": engine.retail_price_source,
        "shop": {
            "total_revenue": round(tax_base, 2),
            "tax_base_revenue": round(tax_base, 2),
            "tax_usn": tax_usn,
            "tax_total": tax_usn,
            "clear_profit": clear_profit,
            "operational_profit": operational_profit,
            "total_sku_margin": engine.total_sku_margin,
            "margin_rate_pct": margin_rate,
            "buyout_coef_pct": round(engine.buyout_coef_pct, 1),
            "drr_pct": drr_pct,
            "total_storage_cost": round(engine.total_storage_cost, 2),
            "total_system_losses": round(engine.total_system_losses, 2),
            "storage_cost": round(engine.total_storage_cost, 2),
            "credit_deductions": round(engine.credit_deductions, 2),
            "penalties_and_other_rub": round(engine.total_system_losses, 2),
            "ad_spend": round(engine.total_ad_spend, 2),
            "logistics_cost": engine.logistics_cost,
            "commission_cost": engine.commission_cost,
            "cost_of_goods": engine.cost_of_goods,
            "sales_qty": round(engine.sales_qty, 2),
            "returns_qty": round(engine.returns_qty, 2),
        },
        "sku_catalog": sku_rows,
        "abc_analysis": {
            "group_A": group_a or ["Лидеры отсутствуют, требуется оптимизация"],
            "group_C": group_c,
            "total_group_c_count": len(group_c),
        },
        "oos_predictions": oos_predictions,
        "year_forecast_rub": round(tax_base * 12, 0),
    }


__all__ = (
    "WB_COLUMN_SYNONYMS",
    "COLUMN_SYNONYMS",
    "CFO_ENGINE_NAME",
    "CfoEngineResult",
    "CfoSkuBucket",
    "aggregate_cfo_engine_v11_1",
    "build_report_metrics_for_history",
    "build_final_metrics_json",
    "build_cfo_metrics_dict_from_rows",
    "sync_table_cfo_processing_worker",
    "check_wb_finance_upload_file",
    "validate_wb_finance_detail_structure",
    "WB_FINANCE_ERROR_INVALID_STRUCTURE",
    "wb_finance_invalid_structure_payload",
    "compute_buyout_coef_pct",
    "find_column_index",
    "resolve_wb_cfo_core_column_indices",
    "should_warn_column_structure",
    "wb_core_finance_columns_recognized",
    "compress_extracted_text",
    "compute_seller_matrix_etl",
    "extract_text_from_document",
    "extract_text_from_pdf",
    "pdf_first_page_to_data_url",
    "render_pdf_first_page_png",
    "read_xlsx_rows_from_bytes",
    "read_xlsx_rows_from_path",
    "SellerMatrixEtl",
    "MatrixAbcSku",
    "MatrixSkuDetail",
    "MatrixOosForecast",
    "DEFAULT_MAX_CHARS",
    "MAX_DOCUMENT_BYTES",
    "MAX_SPREADSHEET_BYTES",
    "SPREADSHEET_SUFFIXES",
    "TXT_DOCUMENT_TOO_BIG",
    "TXT_SPREADSHEET_TOO_BIG",
    "document_too_big_message",
    "is_spreadsheet_suffix",
    "max_document_bytes_for",
    "is_document_size_ok",
    "DocumentTooBigError",
    "download_telegram_document_to_buffer",
    "download_telegram_document_to_path",
)
